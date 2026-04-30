"""
VERSUS — Interface Streamlit
Alignement sémantique de textes.
"""

import streamlit as st
import streamlit.components.v1 as components
from Corpus import Corpus
from Document import Document
from Comparateur import PairText, FAISS_AVAILABLE
from Global_stuff import Global_stuff
from Config import ComparisonConfig, get_document_hash
from guide import render_guide
from sentence_transformers import SentenceTransformer
from collections import Counter
import json
import io
import time
import os

STATE_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), ".vs_state.json")

st.set_page_config(
    page_title="VERSUS",
    page_icon="logo.png",
    layout="wide"
)

st.markdown("""
<style>
    .stApp { background-color: #fafbfc !important; }
    .stApp > header { background-color: #ffffff !important; }
    .stMarkdown, .stText, p, span, label { color: #111827 !important; }

    /* Supprimer le padding Streamlit au-dessus du logo */
    .block-container { padding-top: 0rem !important; padding-bottom: 1rem !important; }
    [data-testid="stToolbar"] { display: none !important; }
    #MainMenu {visibility: hidden;}

    /* Sidebar */
    section[data-testid="stSidebar"] {
        background-color: #f8f9fb !important;
        border-right: 1px solid #e5e7eb !important;
    }
    section[data-testid="stSidebar"] > div { padding-top: 0 !important; }
    [data-testid="collapsedControl"] { display: flex !important; }

    /* Labels sidebar non-cliquables */
    .sidebar-section-title {
        font-size: 0.65rem;
        text-transform: uppercase;
        letter-spacing: 0.1em;
        color: #9ca3af;
        font-weight: 600;
        margin: 1rem 0 0.5rem 0;
    }
    .sidebar-row {
        display: flex;
        justify-content: space-between;
        align-items: center;
        padding: 4px 0;
        border-bottom: 1px solid #f1f5f9;
        font-size: 0.75rem;
    }
    .sidebar-row:last-child { border-bottom: none; }
    .sidebar-key { color: #6b7280; }
    .sidebar-val { font-weight: 600; color: #111827; max-width: 110px; text-align: right; overflow: hidden; text-overflow: ellipsis; white-space: nowrap; }
    .sidebar-badge {
        display: inline-block;
        padding: 1px 7px;
        border-radius: 10px;
        font-size: 0.65rem;
        font-weight: 600;
    }
    .badge-green { background: #d1fae5; color: #065f46; }
    .badge-blue  { background: #dbeafe; color: #1e40af; }
    .badge-gray  { background: #f3f4f6; color: #6b7280; border: 1px solid #e5e7eb; }
    
    .panel-title {
        font-family: 'JetBrains Mono', monospace;
        font-size: 0.7rem;
        text-transform: uppercase;
        letter-spacing: 0.1em;
        color: #4b5563 !important;
        font-weight: 600;
        margin-bottom: 1rem;
    }
    
    .stat-box {
        text-align: center;
        padding: 1rem;
        background: white;
        border-radius: 8px;
        border: 1px solid #e5e7eb;
    }
    .stat-value { font-size: 1.75rem; font-weight: 700; color: #2563eb !important; }
    .stat-label { font-size: 0.65rem; text-transform: uppercase; color: #6b7280 !important; }
    
    .stTextInput input, .stTextArea textarea { background-color: #ffffff !important; color: #111827 !important; }
    .stSelectbox > div > div { background-color: #ffffff !important; color: #111827 !important; }
    
    .stFileUploader > div > button { display: none !important; }
    .stFileUploader section > button { display: none !important; }
    [data-testid="stFileUploader"] button { display: none !important; }
    [data-testid="stFileUploader"] [data-testid="stFileUploaderFile"] { display: none !important; }
  
    button[kind="primary"] p,
    button[kind="primary"] span,
    button[kind="primary"] div,
    .stButton > button[kind="primary"] p,
    .stButton > button[kind="primary"] span {
        color: white !important;
    }
    
    .info-badge {
        display: inline-block;
        padding: 2px 8px;
        border-radius: 4px;
        font-size: 0.7rem;
        font-weight: 600;
        margin-right: 4px;
    }
    .info-badge-green { background: #d1fae5; color: #065f46; }
    .info-badge-blue { background: #dbeafe; color: #1e40af; }
    .info-badge-gray { background: #f3f4f6; color: #374151; }
    
    .stats-table {
        width: 100%;
        border-collapse: collapse;
        font-size: 0.85rem;
    }
    .stats-table th {
        text-align: left;
        padding: 0.5rem 0.75rem;
        background: #f1f5f9;
        color: #475569 !important;
        font-weight: 600;
        border-bottom: 2px solid #e2e8f0;
    }
    .stats-table td {
        padding: 0.4rem 0.75rem;
        border-bottom: 1px solid #e2e8f0;
        color: #1e293b !important;
    }
    .stats-table tr:hover td { background: #f8fafc; }
    .freq-bar {
        display: inline-block;
        height: 8px;
        border-radius: 4px;
        margin-right: 6px;
        vertical-align: middle;
    }
    
    footer {visibility: hidden;}

    /* Masquer la pagination du file_uploader ("Showing page X of Y") */
    [data-testid="stFileUploaderPagination"] {display: none !important;}
</style>
""", unsafe_allow_html=True)


def save_state():
    """Sauvegarde l'état minimal de la session dans un fichier JSON local."""
    try:
        data = {
            "align_th":         st.session_state.get("align_th", 0.85),
            "exclude_stopwords":st.session_state.get("exclude_stopwords", False),
            "active_tab":       st.session_state.get("active_tab", 0),
            "align_sort_mode":  st.session_state.get("align_sort_mode", "Combinée (60/40)"),
            "source": (
                {"name": st.session_state.app_source.name,
                 "content": st.session_state.app_source.text.origin_content}
                if st.session_state.get("app_source") else None
            ),
            "corpus": [
                {"name": doc.name, "content": doc.text.origin_content}
                for doc in st.session_state.app_corpus_full.documents
            ] if "app_corpus_full" in st.session_state else [],
        }
        with open(STATE_FILE, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
    except Exception:
        pass


def load_saved_state():
    """Charge l'état sauvegardé depuis le fichier JSON. Retourne {} si absent."""
    try:
        if os.path.exists(STATE_FILE):
            with open(STATE_FILE, "r", encoding="utf-8") as f:
                return json.load(f)
    except Exception:
        pass
    return {}


def init_state():
    if 'app_config' not in st.session_state:
        saved = load_saved_state()
        cfg = ComparisonConfig()
        st.session_state.app_config = cfg
        st.session_state.app_corpus      = Corpus(config=cfg)
        st.session_state.app_corpus_full = Corpus(config=cfg)
        st.session_state.corpus_filter_kw = ''
        st.session_state.app_rankings    = []
        st.session_state.app_target      = None
        st.session_state.app_alignments  = []
        st.session_state.app_comparateur = None
        st.session_state.active_tab      = saved.get("active_tab", 0)
        st.session_state.align_th        = saved.get("align_th", 0.85)
        st.session_state.app_config.similarity_threshold = st.session_state.align_th
        st.session_state.need_realign    = True
        st.session_state.params_pending  = False
        st.session_state.pending_params  = set()
        st.session_state.ranking_pending = False
        st.session_state.exclude_stopwords = saved.get("exclude_stopwords", False)
        st.session_state.app_config.use_stopwords = st.session_state.exclude_stopwords
        st.session_state.removed_corpus_files = set()
        st.session_state.corpus_uploader_key  = 0
        st.session_state.source_uploader_key  = 0
        st.session_state.align_page      = 0
        st.session_state.align_sort_mode = saved.get("align_sort_mode", "Combinée (60/40)")

        # Restaurer la source
        if saved.get("source"):
            try:
                st.session_state.app_source = Document(
                    name=saved["source"]["name"],
                    content=saved["source"]["content"]
                )
            except Exception:
                st.session_state.app_source = None
        else:
            st.session_state.app_source = None

        # Restaurer le corpus
        for doc_data in saved.get("corpus", []):
            try:
                doc = Document(name=doc_data["name"], content=doc_data["content"])
                st.session_state.app_corpus_full.add_doc(doc)
            except Exception:
                pass
        st.session_state.app_corpus = st.session_state.app_corpus_full


def reset_all():
    try:
        if os.path.exists(STATE_FILE):
            os.remove(STATE_FILE)
    except Exception:
        pass
    st.session_state.app_config = ComparisonConfig()
    st.session_state.app_corpus = Corpus(config=st.session_state.app_config)
    st.session_state.app_corpus_full = Corpus(config=st.session_state.app_config)
    st.session_state.corpus_filter_kw = ''
    st.session_state.app_source = None
    st.session_state.app_rankings = []
    st.session_state.app_target = None
    st.session_state.app_alignments = []
    st.session_state.app_comparateur = None
    st.session_state.active_tab = 0
    st.session_state.align_th = 0.85
    st.session_state.need_realign = True
    st.session_state.params_pending = False
    st.session_state.pending_params = set()
    st.session_state.ranking_pending = False
    st.session_state.exclude_stopwords = False
    st.session_state.removed_corpus_files = set()
    st.session_state.corpus_uploader_key = st.session_state.get('corpus_uploader_key', 0) + 1
    st.session_state.source_uploader_key = st.session_state.get('source_uploader_key', 0) + 1
    st.session_state.align_page = 0
    st.session_state.align_sort_mode = "Combinée (60/40)"
    st.session_state["_cible_display"] = None


@st.cache_resource
def load_model():
    # return SentenceTransformer("sentence-transformers/all-MiniLM-L6-v2")
    return SentenceTransformer("sentence-transformers/paraphrase-MiniLM-L3-v2")


# =================================================================
#  Cache global des embeddings de phrases
# =================================================================

@st.cache_resource
def get_embedding_cache():
    """
    Dictionnaire persistant (niveau processus) pour éviter de recalculer
    les embeddings d'un même document entre deux lancements d'alignement.
    Clé  : MD5 du contenu concaténé des phrases.
    Valeur: ndarray retourné par model.encode().
    """
    return {}


class CachedEncoderModel:
    """
    Wrapper transparent autour de SentenceTransformer.
    Toute appel à .encode(sentences) est mis en cache :
    si les mêmes phrases ont déjà été encodées, le résultat est
    retourné instantanément sans solliciter le GPU/CPU.
    """

    def __init__(self, model):
        self._model = model
        self._cache = get_embedding_cache()

    def encode(self, sentences, **kwargs):
        import hashlib, time

        if hasattr(sentences, "tolist"):
            sentences = sentences.tolist()
        if not isinstance(sentences, (list, tuple)) or len(sentences) == 0:
            return self._model.encode(sentences, **kwargs)

        raw = "|".join(str(s) for s in sentences)
        key = hashlib.md5(raw.encode("utf-8", errors="replace")).hexdigest()

        if key in self._cache:
            print(f"[EmbCache] ✅ HIT  {key[:8]}… {len(sentences)} phrases → 0.000s")
            return self._cache[key]

        t0 = time.time()
        self._cache[key] = self._model.encode(sentences, **kwargs)
        print(f"[EmbCache] 💾 MISS {key[:8]}… {len(sentences)} phrases → {time.time()-t0:.2f}s")
        return self._cache[key]

    def __getattr__(self, name):
        # Délègue tous les autres attributs/méthodes au modèle original
        return getattr(self._model, name)


def read_file(uploaded_file):
    name = uploaded_file.name.lower()
    if name.endswith('.txt'):
        return uploaded_file.getvalue().decode('utf-8', errors='replace')
    elif name.endswith('.docx'):
        try:
            import docx
            doc = docx.Document(io.BytesIO(uploaded_file.getvalue()))
            return '\n'.join([p.text for p in doc.paragraphs])
        except:
            return None
    elif name.endswith('.pdf'):
        try:
            import PyPDF2
            reader = PyPDF2.PdfReader(io.BytesIO(uploaded_file.getvalue()))
            return '\n'.join([p.extract_text() or '' for p in reader.pages])
        except:
            return None
    return None



def go_to_tab(tab_index):
    st.session_state.active_tab = tab_index
    save_state()




# =================================================================
#  Statistiques textuelles
# =================================================================

def compute_text_stats(text_obj):
    """Calcule les statistiques d'un texte."""
    words = [w.content.lower() for w in text_obj.words]
    sentences = text_obj.sentences
    n_words = len(words)
    n_sentences = len(sentences)
    avg_sent_len = n_words / max(n_sentences, 1)
    word_freq = Counter(words)
    unique_words = len(word_freq)
    return {
        "n_words": n_words,
        "n_sentences": n_sentences,
        "avg_sent_len": round(avg_sent_len, 1),
        "unique_words": unique_words,
        "lexical_diversity": round(unique_words / max(n_words, 1), 3),
        "word_freq": word_freq,
    }


def render_stats(stats1, stats2, name1, name2):
    """Affiche les statistiques comparatives."""
    html = f"""
    <table class="stats-table">
        <tr><th></th><th>{name1}</th><th>{name2}</th></tr>
        <tr><td>Nombre de mots</td><td><b>{stats1['n_words']}</b></td><td><b>{stats2['n_words']}</b></td></tr>
        <tr><td>Nombre de phrases</td><td>{stats1['n_sentences']}</td><td>{stats2['n_sentences']}</td></tr>
        <tr><td>Longueur moyenne des phrases (mots)</td><td>{stats1['avg_sent_len']}</td><td>{stats2['avg_sent_len']}</td></tr>
        <tr><td>Mots uniques</td><td>{stats1['unique_words']}</td><td>{stats2['unique_words']}</td></tr>
        <tr><td>Diversité lexicale</td><td>{stats1['lexical_diversity']}</td><td>{stats2['lexical_diversity']}</td></tr>
    </table>
    """
    st.markdown(html, unsafe_allow_html=True)

    common_words = set(stats1['word_freq'].keys()) & set(stats2['word_freq'].keys())
    stopwords = Global_stuff.STOPWORDS
    significant = {w for w in common_words if w not in stopwords and len(w) > 2}

    if significant:
        combined = [(w, stats1['word_freq'][w], stats2['word_freq'][w]) for w in significant]
        combined.sort(key=lambda x: x[1] + x[2], reverse=True)
        top = combined[:10]
        max_freq = max(f1 + f2 for _, f1, f2 in top) if top else 1

        st.markdown("**Fréquences en commun** (hors stopwords, top 10)")
        freq_html = '<table class="stats-table"><tr><th>Mot</th><th>' + name1 + '</th><th>' + name2 + '</th><th></th></tr>'
        for word, f1, f2 in top:
            w1 = int(f1 / max_freq * 100)
            w2 = int(f2 / max_freq * 100)
            freq_html += f'<tr><td><b>{word}</b></td><td>{f1}</td><td>{f2}</td>'
            freq_html += f'<td><span class="freq-bar" style="width:{w1}px;background:#2563eb"></span>'
            freq_html += f'<span class="freq-bar" style="width:{w2}px;background:#f59e0b"></span></td></tr>'
        freq_html += '</table>'
        st.markdown(freq_html, unsafe_allow_html=True)

    only1 = set(stats1['word_freq'].keys()) - set(stats2['word_freq'].keys()) - stopwords
    only2 = set(stats2['word_freq'].keys()) - set(stats1['word_freq'].keys()) - stopwords
    only1_sig = [(w, stats1['word_freq'][w]) for w in only1 if len(w) > 2]
    only2_sig = [(w, stats2['word_freq'][w]) for w in only2 if len(w) > 2]
    only1_sig.sort(key=lambda x: x[1], reverse=True)
    only2_sig.sort(key=lambda x: x[1], reverse=True)

    c1, c2 = st.columns(2)
    with c1:
        if only1_sig:
            st.markdown(f"**Mots propres à {name1}** (top 10)")
            prop_html = '<table class="stats-table"><tr><th>Mot</th><th>Fréq.</th></tr>'
            for w, f in only1_sig[:10]:
                prop_html += f'<tr><td>{w}</td><td>{f}</td></tr>'
            prop_html += '</table>'
            st.markdown(prop_html, unsafe_allow_html=True)
    with c2:
        if only2_sig:
            st.markdown(f"**Mots propres à {name2}** (top 10)")
            prop_html = '<table class="stats-table"><tr><th>Mot</th><th>Fréq.</th></tr>'
            for w, f in only2_sig[:10]:
                prop_html += f'<tr><td>{w}</td><td>{f}</td></tr>'
            prop_html += '</table>'
            st.markdown(prop_html, unsafe_allow_html=True)

# =================================================================
#  Visualisations
# =================================================================

def render_heatmap(alignments, source_text, target_text, name1, name2):
    """Heatmap de densité des alignements (Canvas)."""
    len1 = len(source_text.origin_content)
    len2 = len(target_text.origin_content)
    
    points = []
    for a in alignments:
        pos1, pos2 = a[0], a[1]
        x = ((pos1[0] + pos1[1]) / 2) / max(len1, 1)
        y = ((pos2[0] + pos2[1]) / 2) / max(len2, 1)
        score = float(a[2]) if len(a) > 2 and isinstance(a[2], (int, float)) else 0.5
        seg_len = ((pos1[1] - pos1[0]) + (pos2[1] - pos2[0])) / 2
        r = max(3, min(8, seg_len / 40))
        points.append({"x": round(x, 4), "y": round(y, 4), "s": round(score, 3), "r": round(r, 1)})
    
    points_json = json.dumps(points)
    n1_safe = name1.replace("'", "\\'")
    n2_safe = name2.replace("'", "\\'")
    
    html = """<!DOCTYPE html><html><head><meta charset="UTF-8"><style>
        *{margin:0;padding:0;box-sizing:border-box}
        body{font-family:system-ui,sans-serif;background:#fff}
        .title{font-size:.75rem;font-weight:600;color:#374151;text-transform:uppercase;letter-spacing:.05em;text-align:center;padding:.75rem 0 .5rem}
        canvas{display:block;margin:0 auto}
        .legend{display:flex;justify-content:center;gap:1rem;font-size:.65rem;color:#6b7280;padding:.5rem;flex-wrap:wrap}
        .legend-item{display:flex;align-items:center;gap:4px}
        .legend-dot{width:10px;height:10px;border-radius:50%}
    </style></head><body>
    <div class="title">Carte des alignements</div>
    <canvas id="hm" width="480" height="380"></canvas>
    <div class="legend">
        <div class="legend-item"><div class="legend-dot" style="background:#f59e0b"></div> Score faible</div>
        <div class="legend-item"><div class="legend-dot" style="background:#10b981"></div> Score élevé</div>
        <div class="legend-item" style="color:#94a3b8">┈┈ Diagonale</div>
        <div class="legend-item" style="font-weight:600">__N_PTS__ pts</div>
    </div>
    <script>
    const pts=__POINTS__;
    const cv=document.getElementById('hm');
    const ctx=cv.getContext('2d');
    const W=cv.width,H=cv.height;
    const pad={top:20,right:25,bottom:50,left:60};
    const pw=W-pad.left-pad.right;
    const ph=H-pad.top-pad.bottom;
    
    ctx.fillStyle='#f9fafb';
    ctx.fillRect(pad.left,pad.top,pw,ph);
    
    ctx.strokeStyle='#e5e7eb';ctx.lineWidth=0.5;
    for(let i=0;i<=4;i++){
        let x=pad.left+(pw*i/4),y=pad.top+(ph*i/4);
        ctx.beginPath();ctx.moveTo(x,pad.top);ctx.lineTo(x,pad.top+ph);ctx.stroke();
        ctx.beginPath();ctx.moveTo(pad.left,y);ctx.lineTo(pad.left+pw,y);ctx.stroke();
    }
    
    ctx.strokeStyle='#cbd5e1';ctx.lineWidth=1;ctx.setLineDash([5,5]);
    ctx.beginPath();ctx.moveTo(pad.left,pad.top);ctx.lineTo(pad.left+pw,pad.top+ph);ctx.stroke();
    ctx.setLineDash([]);
    
    for(const p of pts){
        let x=pad.left+p.x*pw,y=pad.top+p.y*ph;
        let t=Math.max(0,Math.min(1,(p.s-0.5)/0.5));
        let r=Math.round(245-t*229),g=Math.round(158+t*27),b=Math.round(11+t*118);
        ctx.fillStyle='rgba('+r+','+g+','+b+',0.7)';
        ctx.beginPath();ctx.arc(x,y,p.r,0,Math.PI*2);ctx.fill();
        ctx.strokeStyle='rgba('+r+','+g+','+b+',0.9)';ctx.lineWidth=0.5;ctx.stroke();
    }
    
    ctx.fillStyle='#6b7280';ctx.font='11px system-ui';ctx.textAlign='center';
    ctx.fillText('__NAME1__ (position →)',pad.left+pw/2,H-8);
    ctx.save();ctx.translate(14,pad.top+ph/2);ctx.rotate(-Math.PI/2);
    ctx.fillText('__NAME2__ (position →)',0,0);ctx.restore();
    
    ctx.fillStyle='#9ca3af';ctx.font='9px system-ui';ctx.textAlign='center';
    for(let i=0;i<=4;i++) ctx.fillText((i*25)+'%',pad.left+pw*i/4,pad.top+ph+15);
    ctx.textAlign='right';
    for(let i=0;i<=4;i++) ctx.fillText((i*25)+'%',pad.left-5,pad.top+ph*i/4+4);
    </script></body></html>"""
    
    html = html.replace('__POINTS__', points_json)
    html = html.replace('__NAME1__', n1_safe)
    html = html.replace('__NAME2__', n2_safe)
    html = html.replace('__N_PTS__', str(len(points)))
    return html


def render_radar(stats1, stats2, name1, name2, stopwords):
    """Radar comparatif des profils textuels (SVG)."""
    import math as _math
    
    # Calcul des métriques supplémentaires
    sw_count1 = sum(f for w, f in stats1['word_freq'].items() if w in stopwords)
    sw_count2 = sum(f for w, f in stats2['word_freq'].items() if w in stopwords)
    density1 = 1 - (sw_count1 / max(stats1['n_words'], 1))
    density2 = 1 - (sw_count2 / max(stats2['n_words'], 1))
    
    hapax1 = sum(1 for w, f in stats1['word_freq'].items() if f == 1) / max(stats1['unique_words'], 1)
    hapax2 = sum(1 for w, f in stats2['word_freq'].items() if f == 1) / max(stats2['unique_words'], 1)
    
    # 6 axes — valeurs brutes
    raw = [
        ("Diversité\nlexicale", stats1['lexical_diversity'], stats2['lexical_diversity']),
        ("Long. moy.\nphrases", stats1['avg_sent_len'], stats2['avg_sent_len']),
        ("Densité\nlexicale", density1, density2),
        ("Hapax", hapax1, hapax2),
        ("Mots\nuniques", stats1['unique_words'], stats2['unique_words']),
        ("Nb\nphrases", stats1['n_sentences'], stats2['n_sentences']),
    ]
    
    # Normalisation 0-1 (par le max des deux)
    axes = []
    for label, v1, v2 in raw:
        mx = max(v1, v2, 0.001)
        axes.append((label, v1 / mx, v2 / mx))
    
    n = len(axes)
    cx, cy, radius = 200, 195, 140
    
    def polar(i, val):
        angle = -_math.pi / 2 + (2 * _math.pi * i / n)
        x = cx + radius * val * _math.cos(angle)
        y = cy + radius * val * _math.sin(angle)
        return round(x, 1), round(y, 1)
    
    # Grille
    grid_svg = ""
    for level in [0.25, 0.5, 0.75, 1.0]:
        pts = " ".join(f"{polar(i, level)[0]},{polar(i, level)[1]}" for i in range(n))
        grid_svg += f'<polygon points="{pts}" fill="none" stroke="#e5e7eb" stroke-width="0.5"/>\n'
    
    # Axes
    axes_svg = ""
    for i in range(n):
        x, y = polar(i, 1.0)
        axes_svg += f'<line x1="{cx}" y1="{cy}" x2="{x}" y2="{y}" stroke="#d1d5db" stroke-width="0.5"/>\n'
    
    # Labels
    labels_svg = ""
    for i, (label, _, _) in enumerate(axes):
        x, y = polar(i, 1.22)
        lines = label.split('\n')
        for li, line in enumerate(lines):
            ly = y + (li - len(lines)/2 + 0.5) * 12
            labels_svg += f'<text x="{x}" y="{ly}" text-anchor="middle" font-size="10" fill="#6b7280" font-family="system-ui">{line}</text>\n'
    
    # Polygones
    pts1 = " ".join(f"{polar(i, axes[i][1])[0]},{polar(i, axes[i][1])[1]}" for i in range(n))
    pts2 = " ".join(f"{polar(i, axes[i][2])[0]},{polar(i, axes[i][2])[1]}" for i in range(n))
    
    # Points sur les sommets
    dots1 = "".join(f'<circle cx="{polar(i, axes[i][1])[0]}" cy="{polar(i, axes[i][1])[1]}" r="3" fill="#2563eb"/>' for i in range(n))
    dots2 = "".join(f'<circle cx="{polar(i, axes[i][2])[0]}" cy="{polar(i, axes[i][2])[1]}" r="3" fill="#f59e0b"/>' for i in range(n))
    
    n1_safe = name1[:20]
    n2_safe = name2[:20]
    
    html = f"""<!DOCTYPE html><html><head><meta charset="UTF-8"><style>
        *{{margin:0;padding:0;box-sizing:border-box}}
        body{{font-family:system-ui,sans-serif;background:#fff}}
        .title{{font-size:.75rem;font-weight:600;color:#374151;text-transform:uppercase;letter-spacing:.05em;text-align:center;padding:.75rem 0 .25rem}}
        .legend{{display:flex;justify-content:center;gap:1.5rem;font-size:.7rem;padding:.5rem}}
        .legend-item{{display:flex;align-items:center;gap:5px}}
        .legend-dot{{width:12px;height:12px;border-radius:3px}}
    </style></head><body>
    <div class="title">Profils textuels comparés</div>
    <svg viewBox="0 0 400 390" width="400" height="390" style="display:block;margin:0 auto">
        {grid_svg}
        {axes_svg}
        <polygon points="{pts1}" fill="rgba(37,99,235,0.15)" stroke="#2563eb" stroke-width="2"/>
        <polygon points="{pts2}" fill="rgba(245,158,11,0.15)" stroke="#f59e0b" stroke-width="2"/>
        {dots1}
        {dots2}
        {labels_svg}
    </svg>
    <div class="legend">
        <div class="legend-item"><div class="legend-dot" style="background:rgba(37,99,235,0.3);border:2px solid #2563eb"></div> <b>{n1_safe}</b></div>
        <div class="legend-item"><div class="legend-dot" style="background:rgba(245,158,11,0.3);border:2px solid #f59e0b"></div> <b>{n2_safe}</b></div>
    </div>
    </body></html>"""
    
    return html


def render_sankey(alignments, source_text, target_text, name1, name2, n_segments=5):
    """Graphe de flux (Sankey) source → cible (SVG)."""
    len1 = len(source_text.origin_content)
    len2 = len(target_text.origin_content)
    
    # Compter les alignements entre segments
    flows = {}
    for a in alignments:
        pos1, pos2 = a[0], a[1]
        mid1 = (pos1[0] + pos1[1]) / 2
        mid2 = (pos2[0] + pos2[1]) / 2
        seg1 = min(int(mid1 / max(len1, 1) * n_segments), n_segments - 1)
        seg2 = min(int(mid2 / max(len2, 1) * n_segments), n_segments - 1)
        key = (seg1, seg2)
        flows[key] = flows.get(key, 0) + 1
    
    if not flows:
        return "<div style='text-align:center;color:#9ca3af;padding:2rem;font-family:system-ui'>Aucun flux à afficher</div>"
    
    max_flow = max(flows.values())
    total = sum(flows.values())
    
    # Dimensions SVG
    W, H = 700, 350
    pad_top, pad_bottom = 45, 20
    bar_w = 22
    left_x = 80
    right_x = W - 80
    usable_h = H - pad_top - pad_bottom
    gap = 6
    bar_h = (usable_h - (n_segments - 1) * gap) / n_segments
    
    # Couleurs des segments source
    colors = ["#2563eb", "#3b82f6", "#60a5fa", "#93c5fd", "#bfdbfe"]
    if n_segments > len(colors):
        colors = colors * ((n_segments // len(colors)) + 1)
    
    svg_bars = ""
    svg_paths = ""
    svg_labels = ""
    
    # Barres source (gauche)
    for i in range(n_segments):
        y = pad_top + i * (bar_h + gap)
        pct_start = round(i * 100 / n_segments)
        pct_end = round((i + 1) * 100 / n_segments)
        svg_bars += f'<rect x="{left_x}" y="{y}" width="{bar_w}" height="{bar_h}" rx="4" fill="{colors[i]}"/>\n'
        svg_labels += f'<text x="{left_x - 8}" y="{y + bar_h / 2 + 4}" text-anchor="end" font-size="10" fill="#6b7280" font-family="system-ui">{pct_start}-{pct_end}%</text>\n'
    
    # Barres cible (droite)
    for j in range(n_segments):
        y = pad_top + j * (bar_h + gap)
        pct_start = round(j * 100 / n_segments)
        pct_end = round((j + 1) * 100 / n_segments)
        svg_bars += f'<rect x="{right_x}" y="{y}" width="{bar_w}" height="{bar_h}" rx="4" fill="#f59e0b" opacity="0.8"/>\n'
        svg_labels += f'<text x="{right_x + bar_w + 8}" y="{y + bar_h / 2 + 4}" text-anchor="start" font-size="10" fill="#6b7280" font-family="system-ui">{pct_start}-{pct_end}%</text>\n'
    
    # Flux (courbes de Bézier)
    for (seg1, seg2), count in sorted(flows.items(), key=lambda x: x[1]):
        y1 = pad_top + seg1 * (bar_h + gap) + bar_h / 2
        y2 = pad_top + seg2 * (bar_h + gap) + bar_h / 2
        x1 = left_x + bar_w
        x2 = right_x
        thickness = max(1.5, min(bar_h * 0.8, (count / max_flow) * bar_h * 0.7))
        opacity = max(0.15, min(0.6, count / max_flow * 0.6))
        
        cp_offset = (x2 - x1) * 0.4
        color = colors[seg1]
        
        svg_paths += f'<path d="M{x1},{y1} C{x1 + cp_offset},{y1} {x2 - cp_offset},{y2} {x2},{y2}" fill="none" stroke="{color}" stroke-width="{round(thickness, 1)}" opacity="{round(opacity, 2)}"/>\n'
        
        # Étiquette au milieu du flux si significatif
        if count >= max(2, total * 0.05):
            mx = (x1 + x2) / 2
            my = (y1 + y2) / 2
            svg_paths += f'<text x="{mx}" y="{my - 2}" text-anchor="middle" font-size="9" fill="#374151" font-weight="600" font-family="system-ui">{count}</text>\n'
    
    n1_safe = name1[:25]
    n2_safe = name2[:25]
    
    html = f"""<!DOCTYPE html><html><head><meta charset="UTF-8"><style>
        *{{margin:0;padding:0;box-sizing:border-box}}
        body{{font-family:system-ui,sans-serif;background:#fff}}
        .title{{font-size:.75rem;font-weight:600;color:#374151;text-transform:uppercase;letter-spacing:.05em;text-align:center;padding:.75rem 0 0}}
        .subtitle{{font-size:.65rem;color:#9ca3af;text-align:center;padding:.25rem 0 .5rem}}
    </style></head><body>
    <div class="title">Flux des alignements</div>
    <div class="subtitle">{total} alignements entre {n_segments} segments</div>
    <svg viewBox="0 0 {W} {H}" width="100%" height="{H}" style="display:block">
        {svg_paths}
        {svg_bars}
        {svg_labels}
        <text x="{left_x + bar_w / 2}" y="{pad_top - 12}" text-anchor="middle" font-size="11" font-weight="700" fill="#2563eb" font-family="system-ui">{n1_safe}</text>
        <text x="{right_x + bar_w / 2}" y="{pad_top - 12}" text-anchor="middle" font-size="11" font-weight="700" fill="#f59e0b" font-family="system-ui">{n2_safe}</text>
    </svg>
    </body></html>"""
    
    return html


def render_timeline(alignments, source_text, target_text, name1, name2):
    """Timeline parallèle : couverture des alignements sur les deux textes."""
    len1 = len(source_text.origin_content)
    len2 = len(target_text.origin_content)
    
    # Collecter les zones alignées
    src_zones = []
    tgt_zones = []
    links = []
    for a in alignments:
        pos1, pos2 = a[0], a[1]
        s1 = pos1[0] / max(len1, 1)
        e1 = pos1[1] / max(len1, 1)
        s2 = pos2[0] / max(len2, 1)
        e2 = pos2[1] / max(len2, 1)
        src_zones.append((s1, e1))
        tgt_zones.append((s2, e2))
        links.append((s1, e1, s2, e2))
    
    # Calculer couverture
    def coverage(zones):
        if not zones:
            return 0
        merged = []
        for s, e in sorted(zones):
            if merged and s <= merged[-1][1]:
                merged[-1] = (merged[-1][0], max(merged[-1][1], e))
            else:
                merged.append((s, e))
        return sum(e - s for s, e in merged)
    
    cov1 = round(coverage(src_zones) * 100, 1)
    cov2 = round(coverage(tgt_zones) * 100, 1)
    
    W, H = 720, 200
    bar_h = 24
    y1 = 55
    y2 = 135
    pad_l, pad_r = 70, 30
    bw = W - pad_l - pad_r
    
    n1_safe = name1[:30].replace("'", "\\'")
    n2_safe = name2[:30].replace("'", "\\'")
    
    # Barres de fond
    svg = f'<rect x="{pad_l}" y="{y1}" width="{bw}" height="{bar_h}" rx="4" fill="#f1f5f9" stroke="#e2e8f0"/>\n'
    svg += f'<rect x="{pad_l}" y="{y2}" width="{bw}" height="{bar_h}" rx="4" fill="#f1f5f9" stroke="#e2e8f0"/>\n'
    
    # Lignes de connexion (d'abord, pour qu'elles soient derrière)
    link_svg = ""
    max_links = min(len(links), 200)  # limiter pour perf
    step = max(1, len(links) // max_links)
    for i in range(0, len(links), step):
        s1, e1, s2, e2 = links[i]
        mx1 = pad_l + ((s1 + e1) / 2) * bw
        mx2 = pad_l + ((s2 + e2) / 2) * bw
        link_svg += f'<line x1="{round(mx1,1)}" y1="{y1 + bar_h}" x2="{round(mx2,1)}" y2="{y2}" stroke="#93c5fd" stroke-width="0.8" opacity="0.3"/>\n'
    
    # Zones source
    for s, e in src_zones:
        x = pad_l + s * bw
        w = max(1, (e - s) * bw)
        svg += f'<rect x="{round(x,1)}" y="{y1}" width="{round(w,1)}" height="{bar_h}" fill="#2563eb" opacity="0.6" rx="2"/>\n'
    
    # Zones cible
    for s, e in tgt_zones:
        x = pad_l + s * bw
        w = max(1, (e - s) * bw)
        svg += f'<rect x="{round(x,1)}" y="{y2}" width="{round(w,1)}" height="{bar_h}" fill="#f59e0b" opacity="0.6" rx="2"/>\n'
    
    # Labels
    labels = f'<text x="{pad_l - 8}" y="{y1 + bar_h/2 + 4}" text-anchor="end" font-size="10" font-weight="600" fill="#2563eb" font-family="system-ui">{n1_safe}</text>\n'
    labels += f'<text x="{pad_l - 8}" y="{y2 + bar_h/2 + 4}" text-anchor="end" font-size="10" font-weight="600" fill="#f59e0b" font-family="system-ui">{n2_safe}</text>\n'
    
    # Graduations
    grads = ""
    for i in range(0, 101, 25):
        x = pad_l + (i / 100) * bw
        grads += f'<text x="{round(x)}" y="{y1 - 8}" text-anchor="middle" font-size="8" fill="#9ca3af" font-family="system-ui">{i}%</text>\n'
    
    html = f"""<!DOCTYPE html><html><head><meta charset="UTF-8"><style>
        *{{margin:0;padding:0;box-sizing:border-box}}
        body{{font-family:system-ui,sans-serif;background:#fff}}
        .title{{font-size:.75rem;font-weight:600;color:#374151;text-transform:uppercase;letter-spacing:.05em;text-align:center;padding:.75rem 0 0}}
        .subtitle{{font-size:.65rem;color:#9ca3af;text-align:center;padding:.25rem 0 .5rem}}
    </style></head><body>
    <div class="title">Couverture des alignements</div>
    <div class="subtitle">Source : {cov1}% couvert · Cible : {cov2}% couvert · {len(alignments)} alignements</div>
    <svg viewBox="0 0 {W} {H}" width="100%" height="{H}" style="display:block">
        {grads}
        {link_svg}
        {svg}
        {labels}
    </svg>
    </body></html>"""
    
    return html


def render_wordcloud(stats1, stats2, name1, name2, stopwords):
    """Nuage de mots différentiel (HTML/CSS)."""
    freq1 = stats1['word_freq']
    freq2 = stats2['word_freq']
    
    # Mots propres à chaque texte
    only1 = {w: f for w, f in freq1.items() if w not in freq2 and w not in stopwords and len(w) > 2}
    only2 = {w: f for w, f in freq2.items() if w not in freq1 and w not in stopwords and len(w) > 2}
    
    # Mots partagés
    shared_words = {w for w in freq1 if w in freq2 and w not in stopwords and len(w) > 2}
    shared = {}
    for w in shared_words:
        f1, f2 = freq1[w], freq2[w]
        shared[w] = (f1, f2, f1 + f2)
    
    def top_n(d, n=25):
        return sorted(d.items(), key=lambda x: x[1] if isinstance(x[1], (int, float)) else x[1][2], reverse=True)[:n]
    
    top_only1 = top_n(only1, 20)
    top_only2 = top_n(only2, 20)
    top_shared = top_n(shared, 25)
    
    # Max freq pour scaling
    all_freqs = [f for _, f in top_only1] + [f for _, f in top_only2] + [t[2] for _, t in top_shared]
    max_f = max(all_freqs) if all_freqs else 1
    
    def size(f, mn=11, mx=28):
        return round(mn + (f / max_f) * (mx - mn), 1)
    
    def make_cloud(items, color_fn):
        spans = []
        for w, f in items:
            freq_val = f if isinstance(f, (int, float)) else f[2]
            s = size(freq_val)
            c = color_fn(w, f)
            spans.append(f'<span style="font-size:{s}px;color:{c};padding:2px 5px;display:inline-block;font-weight:{500 if s < 18 else 700}">{w}</span>')
        return " ".join(spans)
    
    cloud1 = make_cloud(top_only1, lambda w, f: "#2563eb")
    cloud2 = make_cloud(top_only2, lambda w, f: "#d97706")
    
    def shared_color(w, f):
        f1, f2, total = f
        ratio = f1 / max(f1 + f2, 1)
        if ratio > 0.6:
            return "#3b82f6"
        elif ratio < 0.4:
            return "#f59e0b"
        else:
            return "#6b7280"
    
    cloud_shared = make_cloud(top_shared, shared_color)
    
    n1 = name1[:20]
    n2 = name2[:20]
    
    html = f"""<!DOCTYPE html><html><head><meta charset="UTF-8"><style>
        *{{margin:0;padding:0;box-sizing:border-box}}
        body{{font-family:Georgia,serif;background:#fff}}
        .title{{font-size:.75rem;font-weight:600;color:#374151;text-transform:uppercase;letter-spacing:.05em;text-align:center;padding:.75rem 0 .5rem;font-family:system-ui}}
        .grid{{display:grid;grid-template-columns:1fr 1.2fr 1fr;gap:8px;padding:0 12px 12px}}
        .panel{{background:#fafbfc;border:1px solid #e5e7eb;border-radius:8px;padding:10px;text-align:center;min-height:120px;line-height:2}}
        .label{{font-size:.65rem;font-weight:600;text-transform:uppercase;letter-spacing:.05em;margin-bottom:6px;font-family:system-ui}}
        .count{{font-size:.6rem;color:#9ca3af;font-family:system-ui}}
    </style></head><body>
    <div class="title">Vocabulaire différentiel</div>
    <div class="grid">
        <div class="panel">
            <div class="label" style="color:#2563eb">Propres à {n1}</div>
            {cloud1}
            <div class="count">{len(only1)} mots uniques</div>
        </div>
        <div class="panel" style="border-color:#d1d5db">
            <div class="label" style="color:#374151">Partagés</div>
            {cloud_shared}
            <div class="count">{len(shared)} mots · <span style="color:#3b82f6">■</span> {n1} dominant · <span style="color:#f59e0b">■</span> {n2} dominant</div>
        </div>
        <div class="panel">
            <div class="label" style="color:#d97706">Propres à {n2}</div>
            {cloud2}
            <div class="count">{len(only2)} mots uniques</div>
        </div>
    </div>
    </body></html>"""
    
    return html


def render_score_histogram(alignments):
    """Histogramme de distribution des scores de similarité (Canvas)."""
    # Extraire les scores (certains alignements peuvent ne pas avoir de score)
    scores = []
    for a in alignments:
        if len(a) > 2 and isinstance(a[2], (int, float)):
            scores.append(float(a[2]))
    
    if not scores:
        return "<div style='text-align:center;color:#9ca3af;padding:2rem;font-family:system-ui'>Aucun score disponible</div>"
    
    # Bins de 0.05
    bins = {}
    for s in scores:
        b = round(int(s / 0.05) * 0.05, 2)
        b = max(0.5, min(b, 0.95))
        bins[b] = bins.get(b, 0) + 1
    
    all_bins = []
    for i in range(10, 21):
        val = round(i * 0.05, 2)
        all_bins.append((val, bins.get(val, 0)))
    
    max_count = max(c for _, c in all_bins) if all_bins else 1
    avg_score = sum(scores) / len(scores)
    med_score = sorted(scores)[len(scores) // 2]
    
    bins_json = json.dumps(all_bins)
    
    html = f"""<!DOCTYPE html><html><head><meta charset="UTF-8"><style>
        *{{margin:0;padding:0;box-sizing:border-box}}
        body{{font-family:system-ui,sans-serif;background:#fff}}
        .title{{font-size:.75rem;font-weight:600;color:#374151;text-transform:uppercase;letter-spacing:.05em;text-align:center;padding:.75rem 0 0}}
        .subtitle{{font-size:.65rem;color:#9ca3af;text-align:center;padding:.25rem 0 .5rem}}
        canvas{{display:block;margin:0 auto}}
    </style></head><body>
    <div class="title">Distribution des scores</div>
    <div class="subtitle">{len(scores)} scores · Moy: {avg_score:.3f} · Méd: {med_score:.3f}</div>
    <canvas id="hist" width="440" height="250"></canvas>
    <script>
    const bins={bins_json};
    const maxC={max_count};
    const avg={round(avg_score, 4)};
    const cv=document.getElementById('hist');
    const ctx=cv.getContext('2d');
    const W=cv.width,H=cv.height;
    const pad={{top:15,right:20,bottom:45,left:50}};
    const pw=W-pad.left-pad.right;
    const ph=H-pad.top-pad.bottom;
    
    ctx.fillStyle='#f9fafb';
    ctx.fillRect(pad.left,pad.top,pw,ph);
    
    // Grille
    ctx.strokeStyle='#f1f5f9';ctx.lineWidth=0.5;
    for(let i=0;i<=4;i++){{
        let y=pad.top+ph*(1-i/4);
        ctx.beginPath();ctx.moveTo(pad.left,y);ctx.lineTo(pad.left+pw,y);ctx.stroke();
    }}
    
    // Barres
    const bw=pw/bins.length-4;
    for(let i=0;i<bins.length;i++){{
        let [val,count]=bins[i];
        let x=pad.left+i*(bw+4)+2;
        let h=count/maxC*ph;
        let y=pad.top+ph-h;
        
        let t=Math.max(0,Math.min(1,(val-0.5)/0.5));
        let r=Math.round(37+t*0),g=Math.round(99+t*86),b_c=Math.round(235-t*106);
        ctx.fillStyle='rgb('+r+','+g+','+b_c+')';
        ctx.beginPath();
        ctx.roundRect(x,y,bw,h,3);
        ctx.fill();
        
        // Label X
        ctx.fillStyle='#9ca3af';ctx.font='9px system-ui';ctx.textAlign='center';
        ctx.fillText(val.toFixed(2),x+bw/2,pad.top+ph+13);
        
        // Count
        if(count>0){{
            ctx.fillStyle='#374151';ctx.font='bold 9px system-ui';
            ctx.fillText(count,x+bw/2,y-4);
        }}
    }}
    
    // Ligne moyenne
    let avgX=pad.left+((avg-0.5)/0.55)*pw;
    ctx.strokeStyle='#ef4444';ctx.lineWidth=1.5;ctx.setLineDash([4,3]);
    ctx.beginPath();ctx.moveTo(avgX,pad.top);ctx.lineTo(avgX,pad.top+ph);ctx.stroke();
    ctx.setLineDash([]);
    ctx.fillStyle='#ef4444';ctx.font='bold 9px system-ui';ctx.textAlign='center';
    ctx.fillText('moy.',avgX,pad.top+ph+28);
    
    // Axe Y
    ctx.fillStyle='#9ca3af';ctx.font='9px system-ui';ctx.textAlign='right';
    for(let i=0;i<=4;i++){{
        let v=Math.round(maxC*i/4);
        let y=pad.top+ph*(1-i/4);
        ctx.fillText(v,pad.left-5,y+3);
    }}
    </script></body></html>"""
    
    return html


def render_density(alignments, source_text, target_text, name1, name2, n_bins=30):
    """Courbe de densité positionnelle des alignements (Canvas)."""
    len1 = len(source_text.origin_content)
    len2 = len(target_text.origin_content)
    
    density1 = [0] * n_bins
    density2 = [0] * n_bins
    
    for a in alignments:
        pos1, pos2 = a[0], a[1]
        mid1 = (pos1[0] + pos1[1]) / 2
        mid2 = (pos2[0] + pos2[1]) / 2
        b1 = min(int(mid1 / max(len1, 1) * n_bins), n_bins - 1)
        b2 = min(int(mid2 / max(len2, 1) * n_bins), n_bins - 1)
        density1[b1] += 1
        density2[b2] += 1
    
    max_d = max(max(density1), max(density2), 1)
    
    d1_json = json.dumps(density1)
    d2_json = json.dumps(density2)
    n1_safe = name1[:20].replace("'", "\\'")
    n2_safe = name2[:20].replace("'", "\\'")
    
    html = f"""<!DOCTYPE html><html><head><meta charset="UTF-8"><style>
        *{{margin:0;padding:0;box-sizing:border-box}}
        body{{font-family:system-ui,sans-serif;background:#fff}}
        .title{{font-size:.75rem;font-weight:600;color:#374151;text-transform:uppercase;letter-spacing:.05em;text-align:center;padding:.75rem 0 0}}
        .subtitle{{font-size:.65rem;color:#9ca3af;text-align:center;padding:.25rem 0 .5rem}}
        canvas{{display:block;margin:0 auto}}
        .legend{{display:flex;justify-content:center;gap:1.5rem;font-size:.7rem;padding:.5rem}}
        .legend-item{{display:flex;align-items:center;gap:5px}}
        .legend-line{{width:20px;height:3px;border-radius:2px}}
    </style></head><body>
    <div class="title">Densité des alignements par position</div>
    <div class="subtitle">{n_bins} segments · max {max_d} alignements/segment</div>
    <canvas id="dens" width="480" height="230"></canvas>
    <div class="legend">
        <div class="legend-item"><div class="legend-line" style="background:#2563eb"></div> {n1_safe}</div>
        <div class="legend-item"><div class="legend-line" style="background:#f59e0b"></div> {n2_safe}</div>
    </div>
    <script>
    const d1={d1_json},d2={d2_json};
    const maxD={max_d};
    const cv=document.getElementById('dens');
    const ctx=cv.getContext('2d');
    const W=cv.width,H=cv.height;
    const pad={{top:15,right:20,bottom:35,left:45}};
    const pw=W-pad.left-pad.right;
    const ph=H-pad.top-pad.bottom;
    
    ctx.fillStyle='#f9fafb';
    ctx.fillRect(pad.left,pad.top,pw,ph);
    
    // Grille
    ctx.strokeStyle='#f1f5f9';ctx.lineWidth=0.5;
    for(let i=0;i<=4;i++){{
        let y=pad.top+ph*(1-i/4);
        ctx.beginPath();ctx.moveTo(pad.left,y);ctx.lineTo(pad.left+pw,y);ctx.stroke();
    }}
    
    function drawCurve(data,color,fillColor){{
        let n=data.length;
        ctx.beginPath();
        ctx.moveTo(pad.left,pad.top+ph);
        for(let i=0;i<n;i++){{
            let x=pad.left+(i+0.5)/n*pw;
            let y=pad.top+ph*(1-data[i]/maxD);
            if(i===0) ctx.lineTo(x,y);
            else{{
                let px=pad.left+(i-0.5)/n*pw;
                let cpx=(px+x)/2;
                ctx.bezierCurveTo(cpx,pad.top+ph*(1-data[i-1]/maxD),cpx,y,x,y);
            }}
        }}
        ctx.lineTo(pad.left+pw,pad.top+ph);
        ctx.closePath();
        ctx.fillStyle=fillColor;
        ctx.fill();
        
        ctx.beginPath();
        for(let i=0;i<n;i++){{
            let x=pad.left+(i+0.5)/n*pw;
            let y=pad.top+ph*(1-data[i]/maxD);
            if(i===0) ctx.moveTo(x,y);
            else{{
                let px=pad.left+(i-0.5)/n*pw;
                let cpx=(px+x)/2;
                ctx.bezierCurveTo(cpx,pad.top+ph*(1-data[i-1]/maxD),cpx,y,x,y);
            }}
        }}
        ctx.strokeStyle=color;ctx.lineWidth=2.5;ctx.stroke();
    }}
    
    drawCurve(d1,'#2563eb','rgba(37,99,235,0.1)');
    drawCurve(d2,'#f59e0b','rgba(245,158,11,0.1)');
    
    // Axes
    ctx.fillStyle='#9ca3af';ctx.font='9px system-ui';
    ctx.textAlign='center';
    for(let i=0;i<=4;i++) ctx.fillText((i*25)+'%',pad.left+pw*i/4,pad.top+ph+15);
    ctx.textAlign='right';
    for(let i=0;i<=4;i++){{
        let v=Math.round(maxD*i/4);
        ctx.fillText(v,pad.left-5,pad.top+ph*(1-i/4)+4);
    }}
    ctx.textAlign='center';
    ctx.fillText('Position dans le texte (%)',pad.left+pw/2,pad.top+ph+28);
    </script></body></html>"""
    
    return html


def render_cooccurrence(alignments, source_text, target_text, stopwords, n_top=12):
    """Matrice de co-occurrence des mots-clés dans les segments alignés (Canvas)."""
    from collections import Counter as _Counter
    import re as _re
    
    # Extraire les mots significatifs de chaque segment aligné
    word_segments = []
    global_freq = _Counter()
    
    for a in alignments:
        pos1, pos2 = a[0], a[1]
        t1 = source_text.origin_content[pos1[0]:pos1[1]]
        t2 = target_text.origin_content[pos2[0]:pos2[1]]
        combined = t1 + " " + t2
        words = set(w.lower() for w in _re.findall(r'\b\w+\b', combined) if len(w) > 2 and w.lower() not in stopwords)
        word_segments.append(words)
        for w in words:
            global_freq[w] += 1
    
    # Top N mots les plus fréquents
    top_words = [w for w, _ in global_freq.most_common(n_top)]
    n = len(top_words)
    
    if n < 3:
        return "<div style='text-align:center;color:#9ca3af;padding:2rem;font-family:system-ui'>Pas assez de mots-clés pour la matrice</div>"
    
    # Matrice de co-occurrence
    matrix = [[0] * n for _ in range(n)]
    for seg_words in word_segments:
        present = [i for i in range(n) if top_words[i] in seg_words]
        for a in present:
            for b in present:
                if a != b:
                    matrix[a][b] += 1
    
    max_val = max(matrix[i][j] for i in range(n) for j in range(n) if i != j) if n > 1 else 1
    max_val = max(max_val, 1)
    
    matrix_json = json.dumps(matrix)
    words_json = json.dumps(top_words)
    
    cell_size = max(28, min(42, 420 // n))
    W = 100 + n * cell_size + 20
    H = 80 + n * cell_size + 20
    
    html = f"""<!DOCTYPE html><html><head><meta charset="UTF-8"><style>
        *{{margin:0;padding:0;box-sizing:border-box}}
        body{{font-family:system-ui,sans-serif;background:#fff}}
        .title{{font-size:.75rem;font-weight:600;color:#374151;text-transform:uppercase;letter-spacing:.05em;text-align:center;padding:.75rem 0 0}}
        .subtitle{{font-size:.65rem;color:#9ca3af;text-align:center;padding:.25rem 0 .5rem}}
        canvas{{display:block;margin:0 auto}}
    </style></head><body>
    <div class="title">Co-occurrence des mots-clés</div>
    <div class="subtitle">{n} mots les plus fréquents dans les {len(alignments)} segments alignés</div>
    <canvas id="cooc" width="{W}" height="{H}"></canvas>
    <script>
    const m={matrix_json};
    const words={words_json};
    const maxV={max_val};
    const n={n};
    const cs={cell_size};
    const cv=document.getElementById('cooc');
    const ctx=cv.getContext('2d');
    const ox=95,oy=70;
    
    // Cellules
    for(let i=0;i<n;i++){{
        for(let j=0;j<n;j++){{
            let x=ox+j*cs,y=oy+i*cs;
            if(i===j){{
                ctx.fillStyle='#f1f5f9';
            }} else {{
                let t=m[i][j]/maxV;
                let r=Math.round(255-t*218);
                let g=Math.round(255-t*156);
                let b_c=Math.round(255-t*20);
                ctx.fillStyle='rgb('+r+','+g+','+b_c+')';
            }}
            ctx.fillRect(x,y,cs-1,cs-1);
            
            if(i!==j && m[i][j]>0){{
                let t=m[i][j]/maxV;
                ctx.fillStyle=t>0.5?'#fff':'#374151';
                ctx.font='bold '+(cs>35?'11':'9')+'px system-ui';
                ctx.textAlign='center';
                ctx.fillText(m[i][j],x+cs/2-0.5,y+cs/2+3.5);
            }}
        }}
    }}
    
    // Labels
    ctx.fillStyle='#374151';
    ctx.font='10px system-ui';
    ctx.textAlign='right';
    for(let i=0;i<n;i++){{
        let label=words[i].length>10?words[i].substring(0,9)+'…':words[i];
        ctx.fillText(label,ox-5,oy+i*cs+cs/2+3);
    }}
    
    ctx.textAlign='center';
    for(let j=0;j<n;j++){{
        ctx.save();
        ctx.translate(ox+j*cs+cs/2,oy-5);
        ctx.rotate(-Math.PI/4);
        let label=words[j].length>10?words[j].substring(0,9)+'…':words[j];
        ctx.fillText(label,0,0);
        ctx.restore();
    }}
    </script></body></html>"""
    
    return html


# =================================================================
#  Main
# =================================================================

def main():
    init_state()

    # Préchauffage du modèle dès l'ouverture — spinner uniquement au premier lancement
    if not st.session_state.get("_model_ready"):
        load_model()
        st.session_state["_model_ready"] = True

    # === SIDEBAR INFORMATIONNELLE ===
    import base64 as _b64
    _logo_path = os.path.join(os.path.dirname(__file__), "logo.png")
    _sidebar_logo_img = ''
    if os.path.exists(_logo_path):
        with open(_logo_path, "rb") as _f:
            _logo_b64 = _b64.b64encode(_f.read()).decode()
        _sidebar_logo_img = f'<img src="data:image/gif;base64,{_logo_b64}" style="height:1.5rem;vertical-align:middle;margin-right:6px">'

    with st.sidebar:
        # --- Logo cliquable (reset) ---
        components.html(f"""
        <div onclick="
            var btns = window.parent.document.querySelectorAll('button');
            for(var i=0;i<btns.length;i++){{
                if(btns[i].innerText.indexOf('versus_reset')!==-1){{btns[i].click();break;}}
            }}
        " style="display:flex;align-items:center;padding:12px 12px 10px 12px;cursor:pointer;
                 border-bottom:1px solid #e5e7eb;user-select:none;transition:background 0.15s"
           onmouseover="this.style.background='#f8f9fb'" onmouseout="this.style.background='transparent'">
            {_sidebar_logo_img}
            <span style="font-size:1.7rem;font-weight:700;letter-spacing:0.08em;color:#111827;font-family:system-ui,sans-serif">VERSUS</span>
        </div>
        """, height=50)

        # --- État ---
        st.markdown("<div class='sidebar-section-title'>État</div>", unsafe_allow_html=True)

        source = st.session_state.get("app_source")
        source_html = (
            f"<span class='sidebar-badge badge-green'>{source.name[:18]}</span>"
            if source else "<span class='sidebar-badge badge-gray'>—</span>"
        )

        _cible_display = st.session_state.get("_cible_display")
        _target        = st.session_state.get("app_target")

        if _cible_display:
            _badge = "badge-green" if (_target and _target.name == _cible_display) else "badge-blue"
            cible_html = f"<span class='sidebar-badge {_badge}'>{_cible_display[:18]}</span>"
        else:
            cible_docs_list = (
                st.session_state.app_corpus_full.documents
                if "app_corpus_full" in st.session_state else []
            )
            if cible_docs_list:
                cible_names_str = ", ".join(d.name[:12] for d in cible_docs_list)
                if len(cible_names_str) > 20:
                    cible_names_str = cible_names_str[:19] + "…"
                cible_html = f"<span class='sidebar-badge badge-blue'>{cible_names_str}</span>"
            else:
                cible_html = "<span class='sidebar-badge badge-gray'>—</span>"

        st.markdown(f"""
        <div class='sidebar-row'><span class='sidebar-key'>Source</span>{source_html}</div>
        <div class='sidebar-row'><span class='sidebar-key'>Cible</span>{cible_html}</div>
        """, unsafe_allow_html=True)



        # --- Paramètres globaux ---
        st.markdown("<div style='margin-top:0.8rem'></div>", unsafe_allow_html=True)
        st.markdown("<div class='sidebar-section-title'>Paramètres</div>", unsafe_allow_html=True)

        seuil = st.session_state.get("align_th", 0.85)
        sw = st.session_state.get("exclude_stopwords", False)
        sw_html = (
            "<span class='sidebar-badge badge-green'>oui</span>"
            if sw else "<span class='sidebar-badge badge-gray'>non</span>"
        )
        sort_mode = st.session_state.get("align_sort_mode", "Combinée (60/40)")

        st.markdown(f"""
        <div class='sidebar-row'><span class='sidebar-key'>Seuil similarité</span><span class='sidebar-val'>{seuil:.2f}</span></div>
        <div class='sidebar-row'><span class='sidebar-key'>Stopwords exclus</span>{sw_html}</div>
        <div class='sidebar-row'><span class='sidebar-key'>Mode de tri</span><span class='sidebar-val'>{sort_mode}</span></div>
        """, unsafe_allow_html=True)

        # --- Résultats (si alignements disponibles) ---
        alignments = st.session_state.get("app_alignments", [])
        rankings   = st.session_state.get("app_rankings", [])
        target     = st.session_state.get("app_target")
        if alignments and target:
            st.markdown("<div style='margin-top:0.8rem'></div>", unsafe_allow_html=True)
            st.markdown("<div class='sidebar-section-title'>Résultats</div>", unsafe_allow_html=True)

            target_score = None
            for r in rankings:
                if r["doc"].name == target.name:
                    target_score = r["score"]
                    break
            score_html = (
                f"<span class='sidebar-val'>{target_score*100:.1f} %</span>"
                if target_score is not None else "<span class='sidebar-badge badge-gray'>—</span>"
            )

            n_align_res = len(alignments)
            align_res_html = f"<span class='sidebar-badge badge-blue'>{n_align_res}</span>"

            n_identical = sum(
                1 for a in alignments
                if len(a) > 3
                and isinstance(a[2], list) and isinstance(a[3], list)
                and len(a[2]) == 0 and len(a[3]) == 0
            )
            identical_html = f"<span class='sidebar-badge badge-blue'>{n_identical}</span>"

            st.markdown(f"""
            <div class='sidebar-row'><span class='sidebar-key'>Similarité globale</span>{score_html}</div>
            <div class='sidebar-row'><span class='sidebar-key'>Alignements</span>{align_res_html}</div>
            <div class='sidebar-row'><span class='sidebar-key'>Segments identiques</span>{identical_html}</div>
            """, unsafe_allow_html=True)

        st.markdown("""
            <style>
            [data-testid="stSidebar"] > div:first-child {
                padding-bottom: 3rem;
            }
            .sidebar-credit {
                position: fixed;
                bottom: 1rem;
                font-size: 0.68rem;
                color: #b0b7c3;
                line-height: 1.5;
                padding: 0 0.5rem;
            }
            </style>
            <div class="sidebar-credit">
                Développé par Motasem Alrahabi<br>
                ObTIC, Sorbonne Université (2026)
            </div>
        """, unsafe_allow_html=True)
    if st.button("___versus_reset___", key="reset_btn"):
        reset_all()
        st.rerun()
    # Masquer le bouton reset via JS (dans iframe)
    components.html("""
    <script>
    setTimeout(function(){
        var doc = window.parent.document;
        doc.querySelectorAll('button').forEach(function(b){
            if(b.innerText.indexOf('versus_reset') !== -1){
                var el = b.closest('[data-testid="stButton"]') || b.parentElement;
                el.style.cssText = 'height:0;overflow:hidden;margin:0;padding:0;position:absolute';
            }
        });
    }, 200);
    </script>
    """, height=1)

    # FAB haut/bas — injecté via st.markdown directement dans le DOM Streamlit
    # (pas d'iframe : position:fixed fonctionne sur la vraie page)
    st.markdown("""
    <style>
    #vs-top-anchor { position: absolute; top: 0; left: 0; }
    .vs-fab-nav {
        position: fixed; right: 18px; bottom: 24px; z-index: 9999;
        display: flex; flex-direction: column; gap: 3px;
    }
    .vs-fab-nav a {
        display: flex; align-items: center; justify-content: center;
        width: 30px; height: 30px; border-radius: 50%;
        background: #ffffff; border: 1px solid #e5e7eb;
        color: #9ca3af; font-size: 14px; text-decoration: none;
        box-shadow: 0 1px 4px rgba(0,0,0,.10);
        transition: color .15s, border-color .15s;
    }
    .vs-fab-nav a:hover { color: #374151; border-color: #9ca3af; }
    </style>
    <div id="vs-top-anchor"></div>
    <div class="vs-fab-nav">
      <a href="#vs-top-anchor" title="Aller en haut">↑</a>
      <a href="#vs-page-bottom" title="Aller en bas">↓</a>
    </div>
    """, unsafe_allow_html=True)

    st.session_state.app_config.adaptive_threshold = True


    # === NAVIGATION ===
    col1, col2, col3, col4, col5 = st.columns(5)
    with col1:
        if st.button("1 - Chargement", use_container_width=True, 
                    type="primary" if st.session_state.active_tab == 0 else "secondary"):
            go_to_tab(0); st.rerun()
    with col2:
        if st.button("2 - Traitement", use_container_width=True,
                    type="primary" if st.session_state.active_tab == 1 else "secondary"):
            go_to_tab(1); st.rerun()
    with col3:
        if st.button("3 - Résultats", use_container_width=True,
                    type="primary" if st.session_state.active_tab == 2 else "secondary"):
            go_to_tab(2); st.rerun()
    with col4:
        if st.button("Stats et Viz", use_container_width=True,
                    type="primary" if st.session_state.active_tab == 4 else "secondary"):
            go_to_tab(4); st.rerun()
    with col5:
        if st.button("Guide", use_container_width=True,
                    type="primary" if st.session_state.active_tab == 3 else "secondary"):
            go_to_tab(3); st.rerun()
    
    st.divider()

    # ===================== PAGE 1) ENTRÉE =====================
    if st.session_state.active_tab == 0:
        col1, col2 = st.columns(2)
        
        with col1:
            st.markdown('<div class="panel-title">📄 TEXTE SOURCE</div>', unsafe_allow_html=True)
            src_file = st.file_uploader(
                "Glisser-déposer", type=['txt', 'docx', 'pdf'],
                key=f"src_upload_{st.session_state.source_uploader_key}",
                label_visibility="collapsed"
            )
            if src_file:
                # Ne recréer que si c'est un nouveau fichier
                if st.session_state.app_source is None or st.session_state.app_source.name != src_file.name:
                    with st.spinner(f"Chargement de {src_file.name}…"):
                        content = read_file(src_file)
                    if content:
                        st.session_state.app_source = Document(name=src_file.name, content=content)
                        save_state()
            
            if st.session_state.app_source:
                doc = st.session_state.app_source
                c1, c2 = st.columns([5, 1])
                with c1:
                    st.caption(f"📄 {doc.name} ({len(doc.text.words)} mots)")
                with c2:
                    if st.button("✕", key="rm_src"):
                        st.session_state.app_source = None
                        st.session_state.source_uploader_key += 1
                        st.rerun()
        
        with col2:
            st.markdown('<div class="panel-title">📚 CORPUS CIBLE</div>', unsafe_allow_html=True)
            corpus_files = st.file_uploader(
                "Glisser-déposer", type=['txt', 'docx', 'pdf'], 
                accept_multiple_files=True, 
                key=f"corpus_upload_{st.session_state.corpus_uploader_key}", 
                label_visibility="collapsed"
            )
            if corpus_files:
                added = 0
                total_files = len(corpus_files)
                _progress_bar = st.progress(0, text="Chargement des fichiers…") if total_files > 1 else None
                for _fi, f in enumerate(corpus_files):
                    if _progress_bar:
                        _pct = int(_fi / total_files * 100)
                        _progress_bar.progress(_pct, text=f"Chargement : {f.name} ({_fi + 1}/{total_files})")
                    existing_full = st.session_state.app_corpus_full.get_documents_names()
                    file_content = read_file(f)
                    if not file_content:
                        continue
                    if f.name not in existing_full:
                        new_doc = Document(name=f.name, content=file_content)
                        st.session_state.app_corpus_full.add_doc(new_doc)
                        added += 1
                    else:
                        from Config import get_document_hash
                        new_hash = get_document_hash(file_content)
                        idx = existing_full.index(f.name)
                        if new_hash != st.session_state.app_corpus_full.documents[idx].document_hash:
                            st.session_state.app_corpus_full.documents[idx] = Document(name=f.name, content=file_content)
                            st.session_state.app_corpus_full.corpus_weights = None
                            if st.session_state.app_target and st.session_state.app_target.name == f.name:
                                st.session_state.app_target = None
                                st.session_state.app_alignments = []
                            st.session_state.app_rankings = []
                            st.session_state.need_realign = True
                            added += 1
                if _progress_bar:
                    _progress_bar.empty()
                if added > 0:
                    # Re-appliquer le filtre actif s'il y en a un
                    if st.session_state.corpus_filter_kw:
                        st.session_state.app_corpus = st.session_state.app_corpus_full.filter(
                            st.session_state.corpus_filter_kw, ignore_case=True)
                    else:
                        st.session_state.app_corpus = st.session_state.app_corpus_full
                    st.session_state.app_rankings = []
                    st.session_state.need_realign = True
                    save_state()
                    st.rerun()
            
            # Toujours afficher TOUS les fichiers du corpus complet dans le drag & drop
            if len(st.session_state.app_corpus_full) > 0:
                for i, doc in enumerate(st.session_state.app_corpus_full.documents):
                    c1, c2 = st.columns([5, 1])
                    with c1:
                        st.caption(f"📄 {doc.name} ({len(doc.text.words)} mots)")
                    with c2:
                        if st.button("✕", key=f"rm_{i}"):
                            removed_name = st.session_state.app_corpus_full.documents[i].name
                            # Supprimer du corpus complet
                            st.session_state.app_corpus_full = st.session_state.app_corpus_full - i
                            # Re-appliquer le filtre actif si besoin
                            if st.session_state.corpus_filter_kw:
                                st.session_state.app_corpus = st.session_state.app_corpus_full.filter(
                                    st.session_state.corpus_filter_kw, ignore_case=True)
                            else:
                                st.session_state.app_corpus = st.session_state.app_corpus_full
                            st.session_state.corpus_uploader_key += 1
                            st.session_state.app_rankings = []
                            st.session_state.need_realign = True
                            if st.session_state.app_target and st.session_state.app_target.name == removed_name:
                                st.session_state.app_target = None
                                st.session_state.app_alignments = []
                            st.rerun()

        can_classify = (st.session_state.app_source is not None
                        and len(st.session_state.app_corpus_full) > 0)
        if st.button("Valider →", use_container_width=True, disabled=not can_classify, type="primary"):
            # Réinitialiser le filtre mots-clés à chaque retour sur page 2
            st.session_state.corpus_filter_kw = ""
            st.session_state.app_corpus = st.session_state.app_corpus_full
            st.session_state.app_rankings = []
            st.session_state["_cible_display"] = None
            st.session_state.active_tab = 1
            st.rerun()
    
    # ===================== PAGE 2) CLASSEMENT =====================
    elif st.session_state.active_tab == 1:
        if not st.session_state.app_source:
            st.warning("⚠️ Chargez un texte source dans l'onglet Entrée")
        elif len(st.session_state.app_corpus) == 0:
            st.warning("⚠️ Ajoutez des documents au corpus dans l'onglet Entrée")
        else:
            already_ranked = bool(st.session_state.app_rankings)

            # ── Choix du mode de classement ──────────────────────────────────
            st.markdown('<div class="panel-title">MODE DE CLASSEMENT</div>', unsafe_allow_html=True)
            classify_mode = st.radio(
                "Mode",
                ["Par similarité", "Par mots-clés"],
                horizontal=True,
                label_visibility="collapsed",
                key="classify_mode_radio"
            )

            # Stopwords — toujours visible, grisé en mode mots-clés
            sw_kw_mode = (classify_mode == "Par mots-clés")
            sw_help = ("Les stopwords sont des mots grammaticaux très fréquents "
                       "(articles, prépositions, conjonctions…) qui n'apportent pas de sens propre. "
                       "Les exclure concentre la comparaison sur le vocabulaire significatif, ce qui "
                       "améliore la précision pour des textes à contenu dense. À désactiver si la "
                       "structure grammaticale elle-même est pertinente pour l'analyse.")
            if sw_kw_mode:
                sw_help += " (non applicable avec le classement par mots-clés)"
            new_exclude_sw = st.checkbox("Exclure les stopwords",
                                         st.session_state.exclude_stopwords,
                                         key="cb_exclude_sw",
                                         disabled=sw_kw_mode,
                                         help=sw_help)
            if new_exclude_sw != st.session_state.exclude_stopwords:
                st.session_state.exclude_stopwords = new_exclude_sw
                st.session_state.app_config.use_stopwords = new_exclude_sw
                save_state()

            if classify_mode == "Par similarité":
                # Arrêter le filtre mots-clés si on revient en mode similarité
                if st.session_state.get("corpus_filter_kw"):
                    st.session_state.corpus_filter_kw = ""
                    st.session_state.app_corpus = st.session_state.app_corpus_full
                    st.session_state.app_rankings = []

                if st.button("Classer les textes par similarité →", use_container_width=True,
                             type="secondary" if already_ranked else "primary"):
                    with st.spinner("Classement en cours…"):
                        temp = st.session_state.app_corpus.copy()
                        idx = temp.index(st.session_state.app_source)
                        if idx == -1:
                            temp.add_doc(st.session_state.app_source)
                            idx = len(temp) - 1
                        if st.session_state.app_config.use_stopwords:
                            for doc in temp.documents:
                                doc.text.remove_stopwords()
                        result = temp.compare(idx, model=CachedEncoderModel(load_model()), inplace=False)
                        st.session_state.app_rankings = []
                        st.session_state.ranking_pending = False
                        st.session_state.params_pending = False
                        for i, doc in enumerate(result.documents):
                            if doc.name != st.session_state.app_source.name:
                                st.session_state.app_rankings.append({
                                    'doc': doc,
                                    'score': result.similarities[i],
                                    'rank': len(st.session_state.app_rankings) + 1
                                })
                        if st.session_state.app_rankings:
                            st.session_state["_cible_display"] = st.session_state.app_rankings[0]['doc'].name
                    st.rerun()

            else:  # Par mots-clés
                active_kw_filter = st.session_state.get("corpus_filter_kw", "")
                kw_c1, kw_c2, kw_c3 = st.columns([3, 1, 1])
                with kw_c1:
                    st.text_input("Mot-clé", label_visibility="collapsed",
                                  placeholder="Ex: philosophie", key="kw_filter_input")
                with kw_c2:
                    kw_case = st.checkbox("Ignorer la casse", True, key="kw_case_checkbox")
                with kw_c3:
                    kw_val = st.session_state.get("kw_filter_input", "").strip()
                    if st.button("Filtrer", use_container_width=True,
                                 type="primary") and kw_val:
                        with st.spinner(f"Filtrage par « {kw_val} »…"):
                            st.session_state.corpus_filter_kw = kw_val
                            st.session_state.app_corpus = st.session_state.app_corpus_full.filter(
                                kw_val, ignore_case=kw_case)
                        kw_matches = st.session_state.app_corpus.keyword_matches
                        st.session_state.app_rankings = []
                        st.session_state.ranking_pending = False
                        st.session_state.params_pending = False
                        for doc in st.session_state.app_corpus.documents:
                            if doc.name != st.session_state.app_source.name:
                                st.session_state.app_rankings.append({
                                    'doc': doc,
                                    'score': kw_matches.get(doc.name, 0),
                                    'rank': len(st.session_state.app_rankings) + 1
                                })
                        if st.session_state.app_rankings:
                            st.session_state["_cible_display"] = st.session_state.app_rankings[0]['doc'].name
                        st.rerun()

                if active_kw_filter:
                    n_found = len(st.session_state.app_corpus.documents)
                    total_docs = len(st.session_state.app_corpus_full)
                    if n_found == 0:
                        st.warning(f"Aucun document cible ne contient « {active_kw_filter} ».")
                    else:
                        st.markdown(
                            f"<div style='font-size:0.78rem;color:#6b7280;margin:0.4rem 0 0.2rem 0'>"
                            f"<b style='color:#2563eb'>{n_found}</b> document(s) retenu(s) sur "
                            f"<b>{total_docs}</b> — filtre : <b>« {active_kw_filter} »</b></div>",
                            unsafe_allow_html=True
                        )
            
            if st.session_state.app_rankings:
                doc_names = [r['doc'].name for r in st.session_state.app_rankings]
                
                source_name = st.session_state.app_source.name
                st.markdown(f'Sélectionner un document à comparer avec : **"{source_name}"**')
                selected_doc = st.selectbox("", doc_names, key="select_align_doc",
                                            label_visibility="collapsed")
                # Toujours synchroniser _cible_display avec le choix courant du selectbox
                st.session_state["_cible_display"] = selected_doc

                # Paramètre seuil cosinus (déplacé depuis la sidebar)
                st.markdown("Seuil cosinus", help="Score minimum de similarité cosinus qu'une paire de segments doit atteindre pour être retenue. C'est le paramètre qui a le plus d'impact sur le volume de résultats. La comparaison des embeddings est faite par cosinus (ou via FAISS si le texte cible dépasse 50 phrases).\n\nUn seuil élevé (> 0,85) ne garde que les correspondances très proches ; un seuil bas (< 0,70) détecte plus de rapprochements mais introduit davantage de faux positifs.")
                new_th = st.slider(
                    "Seuil cosinus", 0.5, 1.0, st.session_state.align_th, 0.01,
                    key="slider_th",
                    label_visibility="collapsed"
                )
                if abs(new_th - st.session_state.align_th) > 0.001:
                    st.session_state.align_th = new_th
                    save_state()

                if st.button("Aligner les textes →", use_container_width=True, type="primary"):
                    for r in st.session_state.app_rankings:
                        if r['doc'].name == selected_doc:
                            st.session_state.app_target = r['doc']
                            st.session_state["_cible_display"] = r['doc'].name
                            break
                    st.session_state.app_alignments = []
                    st.session_state.need_realign = True
                    st.session_state.active_tab = 2
                    st.rerun()
    
    # ===================== PAGE 3) RÉSULTATS =====================
    elif st.session_state.active_tab == 2:
        if not st.session_state.app_source:
            st.warning("⚠️ Chargez un texte source dans l'onglet Entrée")
        elif not st.session_state.app_target:
            st.info("ℹ️ Sélectionnez un document depuis l'onglet Classement")
        else:
            cfg = st.session_state.app_config
            
            th = st.session_state.align_th
            
            # Si params_pending, attendre que l'utilisateur relance manuellement
            # sauf si need_realign est True (premier calcul ou nouveau document)
            if st.session_state.params_pending and not st.session_state.need_realign and st.session_state.app_alignments:
                pass  # warning already shown in sidebar, wait for user action
            
            if st.session_state.need_realign:
                st.session_state.need_realign = False
                st.session_state.params_pending = False
                st.session_state.align_page = 0
                
                progress_bar = st.progress(0, text="Préparation...")
                start_time = time.time()
                
                progress_bar.progress(10, text="Chargement du modèle...")
                model = CachedEncoderModel(load_model())
                
                progress_bar.progress(20, text="Configuration...")
                run_cfg = ComparisonConfig(**st.session_state.app_config.to_dict())
                run_cfg.similarity_threshold = th
                
                progress_bar.progress(30, text="Segmentation en phrases...")
                
                text1 = st.session_state.app_source.text
                text2 = st.session_state.app_target.text
                text1.default()
                text2.default()
                
                if st.session_state.exclude_stopwords:
                    text1.remove_stopwords()
                    text2.remove_stopwords()
                
                comp = PairText(text1, text2, config=run_cfg)
                
                progress_bar.progress(50, text="Alignement par phrases...")
                
                st.session_state.app_alignments = comp.compare_n_grams(
                    model=model, score_threshold=th
                )
                
                progress_bar.progress(80, text="Calcul suppressions/insertions...")
                st.session_state.app_alignments = comp.compute_diffs(
                    st.session_state.app_alignments
                )
                
                st.session_state.app_comparateur = comp
                st.session_state.params_pending = False
                st.session_state.pending_params = set()
                
                elapsed = time.time() - start_time
                progress_bar.progress(100, text=f"✓ Terminé en {elapsed:.1f}s")
                time.sleep(0.3)
                progress_bar.empty()


            # Résultats
            if st.session_state.app_alignments:

                if st.session_state.app_comparateur:
                    # Stopwords pour filtrage mots communs côté JS
                    stopwords = Global_stuff.STOPWORDS

                    # ── Tri dans la page Résultats ──
                    sort_options = ["Lexicale (Jaccard)", "Lexicale (Overlap)", "Sémantique", "Combinée (60/40)"]
                    cur_sort = st.session_state.get('align_sort_mode', 'Combinée (60/40)')
                    sort_choice = st.radio(
                        "Trier les résultats par similarité",
                        sort_options,
                        index=sort_options.index(cur_sort) if cur_sort in sort_options else 3,
                        horizontal=True,
                        key="results_sort_mode",
                        help="**Lexicale (Jaccard)** : ratio mots en commun / union des mots (hors stopwords). Sensible à la longueur des segments.\n\n**Lexicale (Overlap)** : ratio mots en commun / plus petit des deux segments (hors stopwords). Corrige le biais de longueur de Jaccard.\n\n**Sémantique** : score cosinus entre les embeddings. Mesure la proximité de sens indépendamment de la forme lexicale.\n\n**Combinée (60/40)** : 60% sémantique + 20% Jaccard + 20% Overlap, après normalisation min-max."
                    )
                    if sort_choice != st.session_state.align_sort_mode:
                        st.session_state.align_sort_mode = sort_choice
                        st.session_state.align_page = 0
                        st.rerun()

                    mode_map = {
                        "Lexicale (Jaccard)": "lexical_jaccard",
                        "Lexicale (Overlap)": "lexical_overlap",
                        "Sémantique":         "semantic",
                        "Combinée (60/40)":   "combined",
                    }
                    sorted_alignments = st.session_state.app_comparateur.sort_alignments(
                        st.session_state.app_alignments,
                        mode=mode_map.get(sort_choice, "combined")
                    )

                    # ── Pagination ──
                    BLOCKS_PER_PAGE = 30
                    all_matches = sorted_alignments
                    total_blocks = len(all_matches)
                    total_pages = max(1, (total_blocks + BLOCKS_PER_PAGE - 1) // BLOCKS_PER_PAGE)
                    
                    # Clamp page
                    if st.session_state.align_page >= total_pages:
                        st.session_state.align_page = total_pages - 1
                    if st.session_state.align_page < 0:
                        st.session_state.align_page = 0
                    
                    page = st.session_state.align_page
                    start_idx = page * BLOCKS_PER_PAGE
                    end_idx = min(start_idx + BLOCKS_PER_PAGE, total_blocks)
                    page_matches_raw = all_matches[start_idx:end_idx]
                    
                    html = make_html(
                        page_matches_raw,
                        st.session_state.app_comparateur,
                        stopwords,
                        start_offset=start_idx,
                        total_alignments=total_blocks,
                        raw_scores=getattr(st.session_state.app_comparateur, '_raw_scores', {}),
                        all_matches=sorted_alignments
                    )
                    components.html(html, height=1650, scrolling=True)
                    
                    # Navigation si plus d'une page
                    if total_pages > 1:
                        _c1, _c2, _c3 = st.columns([1, 3, 1])
                        with _c1:
                            if st.button("«‹ Préc.", disabled=(page == 0), key="pg_prev", use_container_width=True):
                                st.session_state.align_page -= 1; st.rerun()
                        with _c2:
                            st.markdown(
                                f"<div style='text-align:center;font-size:.75rem;color:#6b7280;"
                                f"border:1px solid #e5e7eb;border-radius:6px;padding:0 14px;height:38px;"
                                f"display:flex;align-items:center;justify-content:center;gap:5px'>"
                                f"<b style='color:#111827'>{page+1}</b>"
                                f"<span style='color:#d1d5db'>/</span><span>{total_pages}</span>"
                                f"<span style='color:#e5e7eb;margin:0 4px'>·</span>"
                                f"<span>blocs {start_idx+1}–{end_idx}"
                                f"<span style='color:#9ca3af'> / {total_blocks}</span></span></div>",
                                unsafe_allow_html=True
                            )
                        with _c3:
                            if st.button("Suiv. ›»", disabled=(page >= total_pages - 1), key="pg_next", use_container_width=True):
                                st.session_state.align_page += 1; st.rerun()

                    # Ancre bas de page (cible du bouton ↓)
                    st.markdown('<div id="vs-page-bottom"></div>', unsafe_allow_html=True)
                

            else:
                st.info("Aucun alignement trouvé avec ces paramètres. Essayez de baisser le seuil.")
    
    # ===================== PAGE 4) STATS ET VIZ =====================
    elif st.session_state.active_tab == 4:
        if not hasattr(st.session_state, 'app_source') or st.session_state.app_source is None:
            st.warning("⚠️ Chargez un texte source dans l'onglet Entrée")
        elif not hasattr(st.session_state, 'app_alignments') or not st.session_state.app_alignments:
            st.info("ℹ️ Lancez un alignement dans l'onglet Résultats pour afficher les statistiques et visualisations.")
        else:
            stats1 = compute_text_stats(st.session_state.app_source.text)
            stats2 = compute_text_stats(st.session_state.app_target.text)

            # ── Visualisations ──
            st.markdown('<div style="font-size:1.1rem;font-weight:700;text-transform:uppercase;letter-spacing:0.08em;color:#111827;text-align:center;margin:1.5rem 0 1rem">📈 Visualisations</div>', unsafe_allow_html=True)

            viz_c1, viz_c2 = st.columns(2)
            with viz_c1:
                heatmap_html = render_heatmap(
                    st.session_state.app_alignments,
                    st.session_state.app_source.text,
                    st.session_state.app_target.text,
                    st.session_state.app_source.name,
                    st.session_state.app_target.name
                )
                components.html(heatmap_html, height=480)
                st.caption("Chaque point représente un passage aligné : sa position horizontale indique où il se situe dans le texte source, sa position verticale dans le texte cible. La couleur varie du jaune (score faible) au vert (score élevé). Une diagonale bien marquée révèle un emprunt linéaire ; des points dispersés signalent des réorganisations structurelles.")
            with viz_c2:
                radar_html = render_radar(
                    stats1, stats2,
                    st.session_state.app_source.name,
                    st.session_state.app_target.name,
                    Global_stuff.STOPWORDS
                )
                components.html(radar_html, height=480)
                st.caption("Comparaison des profils stylistiques sur six dimensions : diversité lexicale, longueur moyenne des phrases, densité lexicale (proportion de mots porteurs de sens), hapax (mots n’apparaissant qu’une fois), nombre de mots uniques et nombre de phrases. Plus le polygone d’un texte est étendu sur un axe, plus il surpasse l’autre sur cette dimension.")

            sankey_html = render_sankey(
                st.session_state.app_alignments,
                st.session_state.app_source.text,
                st.session_state.app_target.text,
                st.session_state.app_source.name,
                st.session_state.app_target.name
            )
            components.html(sankey_html, height=400)
            st.caption("Les deux textes sont découpés en segments de taille égale (colonnes gauche et droite). Chaque ruban relie un segment source au segment cible avec lequel il présente le plus d’alignements ; la largeur du ruban est proportionnelle au nombre de correspondances détectées. Ce graphique révèle d’éventuelles réorganisations de blocs entre les deux textes.")

            # ── Timeline parallèle ──
            timeline_html = render_timeline(
                st.session_state.app_alignments,
                st.session_state.app_source.text,
                st.session_state.app_target.text,
                st.session_state.app_source.name,
                st.session_state.app_target.name
            )
            components.html(timeline_html, height=240)
            st.caption("Les barres représentent la longueur de chaque texte (de gauche à droite = du début à la fin). Les portions colorées indiquent les passages couverts par au moins un alignement ; les traits bleus pâles relient les passages correspondants entre les deux textes. Le pourcentage de couverture affiché mesure la proportion du texte effectivement prise dans des alignements.")

            # ── Histogramme + Densité côte à côte ──
            viz_c3, viz_c4 = st.columns(2)
            with viz_c3:
                hist_html = render_score_histogram(st.session_state.app_alignments)
                components.html(hist_html, height=310)
                st.caption("Répartition des scores de similarité combinés par intervalles de 0,05. Un pic vers la droite indique une majorité d’alignements de haute qualité ; une distribution étalée suggère des correspondances de niveaux variés. La ligne rouge en pointillés marque le score moyen.")
            with viz_c4:
                density_html = render_density(
                    st.session_state.app_alignments,
                    st.session_state.app_source.text,
                    st.session_state.app_target.text,
                    st.session_state.app_source.name,
                    st.session_state.app_target.name
                )
                components.html(density_html, height=310)
                st.caption("Densité des alignements selon leur position dans chaque texte : chaque texte est découpé en tranches et la courbe indique combien d’alignements tombent dans chaque tranche. Les pics révèlent les zones de forte concentration d’emprunts ; les creux signalent des passages peu ou pas repris.")

            # ── Nuage de mots différentiel ──
            wordcloud_html = render_wordcloud(
                stats1, stats2,
                st.session_state.app_source.name,
                st.session_state.app_target.name,
                Global_stuff.STOPWORDS
            )
            components.html(wordcloud_html, height=300)
            st.caption("Vocabulaire exclusif à chaque texte (mots absents de l’autre, hors stopwords). La taille de chaque mot est proportionnelle à sa fréquence. Ce graphique met en évidence les champs lexicaux propres à chaque document et les termes que l’un a ajoutés ou supprimés par rapport à l’autre.")

            # ── Matrice de co-occurrence ──
            cooc_html = render_cooccurrence(
                st.session_state.app_alignments,
                st.session_state.app_source.text,
                st.session_state.app_target.text,
                Global_stuff.STOPWORDS
            )
            components.html(cooc_html, height=520)
            st.caption("Réseau des mots qui apparaissent fréquemment dans les mêmes passages alignés. Chaque nœud représente un terme (hors stopwords) ; chaque lien indique que les deux mots co-occurrent dans au moins un alignement commun. L’épaisseur du lien est proportionnelle à la force de cette co-occurrence. Ce graphique révèle les thèmes et les associations lexicales partagés entre les deux textes.")

            # ── Statistiques ──
            st.divider()
            st.markdown('<div style="font-size:1.1rem;font-weight:700;text-transform:uppercase;letter-spacing:0.08em;color:#111827;text-align:center;margin:1.5rem 0 1rem">📊 Statistiques</div>', unsafe_allow_html=True)
            render_stats(
                stats1, stats2,
                st.session_state.app_source.name,
                st.session_state.app_target.name
            )

    # ===================== PAGE 5) GUIDE =====================
    elif st.session_state.active_tab == 3:
        render_guide()

# =================================================================
#  Rendu HTML des alignements
#  Rendu HTML des alignements
# =================================================================

def make_html(matches, comp, stopwords, start_offset=0, total_alignments=None, raw_scores=None, all_matches=None):
    ctx = 50
    
    COLOR_MATCH = "rgba(16,185,129,0.35)"
    COLOR_SUPPRESS = "rgba(239,68,68,0.4)"
    COLOR_INSERT = "rgba(59,130,246,0.4)"
    BORDER_COLOR = "#10b981"

    import re as _re

    html = """<!DOCTYPE html><html><head><meta charset="UTF-8">
    <style>
        *{box-sizing:border-box;margin:0;padding:0}
        body{font-family:system-ui,sans-serif;background:#fafbfc;padding:1rem;font-size:13px;color:#111827}
        .m{background:#fff;border-radius:12px;margin-bottom:1rem;overflow:hidden;position:relative}
        .mh{padding:.75rem 1rem;font-weight:600;color:#374151;display:flex;justify-content:space-between;align-items:center;background:rgba(16,185,129,.12);border-bottom:1px solid #10b981}
        .mb{display:grid;grid-template-columns:1fr 1fr}
        .p{padding:1rem;border-right:1px solid #e5e7eb}
        .p:last-child{border-right:none}
        .ph{font-size:.7rem;text-transform:uppercase;color:#6b7280;margin-bottom:.5rem;font-weight:600}
        .pn{font-family:Georgia,serif;font-size:1rem;font-weight:600;margin-bottom:.75rem;color:#111827}
        .t{font-family:Georgia,serif;line-height:1.9;background:#f9fafb;padding:1rem;border-radius:6px;max-height:180px;overflow-y:auto;color:#111827}
        .sup{background:rgba(239,68,68,0.3);padding:1px 2px;border-radius:2px;font-weight:600}
        .ins{background:rgba(59,130,246,0.3);padding:1px 2px;border-radius:2px;font-weight:600}
        .cw{font-weight:900;text-decoration:underline;text-underline-offset:3px;font-size:1.15em;letter-spacing:0.03em;text-shadow:0 0 0.5px rgba(0,0,0,0.3);background:rgba(0,0,0,0.06);padding:0 2px;border-radius:2px}
        .ps{font-size:.7rem;color:#6b7280;margin-top:.5rem}
        .bt{padding:.4rem .75rem;background:#e5e7eb;border:none;border-radius:6px;cursor:pointer;font-size:.7rem;margin:0 2px;color:#374151}
        .bt:hover{background:#d1d5db}
        .cnt{font-size:.65rem;padding:3px 8px;border-radius:4px;font-weight:600;color:#374151;background:#f1f5f9;margin-left:.5rem}
        .ctrls{display:flex;align-items:center;gap:1rem;padding:.5rem 1rem;background:#f8fafc;border-bottom:1px solid #e2e8f0;font-size:.75rem;flex-wrap:wrap}
        .ctrls label{display:flex;align-items:center;gap:4px;cursor:pointer;user-select:none;color:#374151}
        .ctrls input[type=checkbox]{accent-color:#10b981;width:14px;height:14px}
        .lbl-sup{color:#dc2626;font-weight:600}
        .lbl-ins{color:#2563eb;font-weight:600}
        .lbl-cw{color:#059669;font-weight:600}
        .lbl-sel{color:#6b7280;font-weight:500}
        .sel-check{accent-color:#2563eb !important}
        .toolbar{display:flex;align-items:center;justify-content:space-between;padding:.75rem 1rem;background:#fff;border:1px solid #e5e7eb;border-radius:8px;margin-bottom:1rem;position:relative}
        .toolbar-left{display:flex;align-items:center;gap:.75rem;font-size:.8rem;color:#374151}
        .toolbar-center{position:absolute;left:50%;transform:translateX(-50%);font-size:1.15rem;font-weight:700;color:#111827;white-space:nowrap}
        .export-btn{padding:.4rem .75rem;background:#f3f4f6;border:1px solid #d1d5db;border-radius:6px;cursor:pointer;font-size:.72rem;font-weight:600;color:#374151}
        .export-btn:hover{background:#e5e7eb}
        .export-btn:disabled{background:#f9fafb;color:#9ca3af;border-color:#e5e7eb;cursor:not-allowed}
        .sel-all-btn{padding:.4rem .75rem;background:#f3f4f6;border:1px solid #d1d5db;border-radius:6px;cursor:pointer;font-size:.72rem;color:#374151}
        .sel-all-btn:hover{background:#e5e7eb}
        .lex-bar-wrap{display:inline-flex;align-items:center;gap:5px;margin-left:.6rem;vertical-align:middle}
        .lex-bar-bg{width:48px;height:6px;background:#e5e7eb;border-radius:3px;overflow:hidden;display:inline-block;vertical-align:middle;border:1px solid #9ca3af}
        .lex-bar-fg{height:100%;border-radius:3px}
        .lex-pct{font-size:.6rem;font-weight:600;min-width:26px;text-align:right}
        .score-bars{display:inline-flex;align-items:center;gap:8px;margin-left:.75rem;vertical-align:middle}
        .score-item{display:inline-flex;align-items:center;gap:4px}
        .score-lbl{font-size:.58rem;font-weight:600;text-transform:uppercase;letter-spacing:.03em;white-space:nowrap}
        .score-bar-bg{width:44px;height:6px;background:#e5e7eb;border-radius:3px;overflow:hidden;display:inline-block;border:1px solid #9ca3af}
        .score-bar-fg{height:100%;border-radius:3px;display:block}
        .score-pct{font-size:.6rem;font-weight:600;min-width:24px}
    </style></head><body>"""
    
    # Toolbar with select all + export
    html += f"""
    <div class="toolbar">
        <div class="toolbar-left">
            <button class="sel-all-btn" id="sel-all-btn" onclick="toggleAll()">☑ Tout sélectionner</button>
            <span id="sel-count">0 sélectionné(s)</span>
        </div>
        <div class="toolbar-center">{total_alignments} alignement(s) au total</div>
        <button class="export-btn" id="export-btn" onclick="exportCSV()">↓ Export CSV</button>
    </div>"""
    
    for i, match in enumerate(matches):
        pos1, pos2 = match[0], match[1]
        n_supp = len(match[2]) if len(match) > 2 and isinstance(match[2], list) and match[2] else 0
        n_ins = len(match[3]) if len(match) > 3 and isinstance(match[3], list) and match[3] else 0

        counts = []
        if n_supp > 0:
            counts.append(f'{n_supp} supp.')
        if n_ins > 0:
            counts.append(f'{n_ins} ins.')
        count_str = f'<span class="cnt">{" · ".join(counts)}</span>' if counts else ''

        # Récupération des scores depuis _raw_scores
        rs = (raw_scores or {}).get((pos1, pos2), {})
        sem_raw  = rs.get('semantic',        0.0)
        jac_raw  = rs.get('lexical_jaccard', 0.0)
        ovl_raw  = rs.get('lexical_overlap', 0.0)
        comb_raw = rs.get('combined',        0.0)

        sem_pct  = int(round(sem_raw  * 100))
        jac_pct  = int(round(jac_raw  * 100))
        ovl_pct  = int(round(ovl_raw  * 100))
        comb_pct = int(round(comb_raw * 100))

        def lex_color(p):
            if p >= 40: return "#10b981", "#065f46"
            if p >= 15: return "#f59e0b", "#92400e"
            return "#ef4444", "#991b1b"

        jc, jpc = lex_color(jac_pct)
        oc, opc = lex_color(ovl_pct)

        def bar(label, pct, bar_col, pct_col, title):
            return (
                f'<span class="score-item" title="{title}">'
                f'<span class="score-lbl" style="color:{pct_col}">{label}</span>'
                f'<span class="score-bar-bg"><span class="score-bar-fg" style="width:{pct}%;background:{bar_col}"></span></span>'
                f'<span class="score-pct" style="color:{pct_col}">{pct}%</span>'
                f'</span>'
            )

        scores_badge = (
            f'<span class="score-bars">'
            + bar("jac.", jac_pct, "#628990", "#374151", "")
            + bar("ovl.", ovl_pct, "#628990", "#374151", "")
            + bar("sém.", sem_pct, "#628990", "#374151", "")
            + bar("cmb.", comb_pct,"#628990", "#374151", "")
            + f'</span>'
        )
        
        html += f"""
        <div class="m" style="border:2px solid {BORDER_COLOR}" id="block_{i}">
            <div class="mh">
                <span>#{start_offset+i+1}{count_str}{scores_badge}</span>
                <div><button class="bt" onclick="c({i},-100)">−</button><button class="bt" onclick="r({i})">Reset</button><button class="bt" onclick="c({i},100)">+</button></div>
            </div>
            <div class="ctrls">
                <label><input type="checkbox" id="chk_sel_{i}" class="sel-check" onchange="onCheckChange({i})"><span class="lbl-sel">Sélectionner</span></label>
                <span style="color:#d1d5db">│</span>
                <label><input type="checkbox" id="chk_sup_{i}" onchange="u({i})"><span class="lbl-sup">Suppressions</span></label>
                <label><input type="checkbox" id="chk_ins_{i}" onchange="u({i})"><span class="lbl-ins">Insertions</span></label>
                <label><input type="checkbox" id="chk_cw_{i}" onchange="u({i})"><span class="lbl-cw">Mots communs</span></label>
            </div>
            <div class="mb">
                <div class="p"><div class="ph">Texte Source</div><div class="pn">{comp.text1.name}</div><div class="t" id="a{i}"></div><div class="ps">Pos: {pos1[0]}-{pos1[1]}</div></div>
                <div class="p"><div class="ph">Texte Cible</div><div class="pn">{comp.text2.name}</div><div class="t" id="b{i}"></div><div class="ps">Pos: {pos2[0]}-{pos2[1]}</div></div>
            </div>
        </div>"""
    
    t1 = comp.text1.origin_content.replace('\\','\\\\').replace('`','\\`').replace('$','\\$').replace('\r',' ').replace('\n',' ')
    t2 = comp.text2.origin_content.replace('\\','\\\\').replace('`','\\`').replace('$','\\$').replace('\r',' ').replace('\n',' ')
    
    mjs = []
    for m_item in matches:
        mjs.append([
            list(m_item[0]), list(m_item[1]),
            list(m_item[2]) if len(m_item) > 2 and isinstance(m_item[2], list) and m_item[2] else [],
            list(m_item[3]) if len(m_item) > 3 and isinstance(m_item[3], list) and m_item[3] else []
        ])
    
    # Tableau compact de TOUS les alignements (toutes pages) pour la sélection globale
    # Format léger : [pos1, pos2, n_supp, n_ins]
    _all = all_matches if all_matches is not None else matches
    all_mjs = []
    for m_item in _all:
        n_s = len(m_item[2]) if len(m_item) > 2 and isinstance(m_item[2], list) else 0
        n_i = len(m_item[3]) if len(m_item) > 3 and isinstance(m_item[3], list) else 0
        all_mjs.append([list(m_item[0]), list(m_item[1]), n_s, n_i])

    # Escape stopwords for JS
    sw_escaped = json.dumps(sorted(list(stopwords)))

    html += f"""
    <script>
    const t1=`{t1}`,t2=`{t2}`,m={json.dumps(mjs)};
    const startOffset={start_offset};
    const BG='{COLOR_MATCH}',SUP='{COLOR_SUPPRESS}',INS='{COLOR_INSERT}';
    const stopWords=new Set({sw_escaped});
    let cs=new Array(m.length).fill({ctx});

    // ── Sélection globale cross-pages ──────────────────────────────────────
    const allM={json.dumps(all_mjs)};
    let selectedGlobal=new Set();
    let allPagesSelected=false;

    function e(t){{let d=document.createElement('div');d.textContent=t;return d.innerHTML}}

    // --- MOTS COMMUNS PAR BLOC ---
    function getBlockCW(i){{
        let p1=m[i][0],p2=m[i][1];
        let context = cs[i];
        let srcFull = t1.substring(Math.max(0, p1[0] - context), Math.min(t1.length, p1[1] + context)).toLowerCase();
        let tgtFull = t2.substring(Math.max(0, p2[0] - context), Math.min(t2.length, p2[1] + context)).toLowerCase();
        let w1 = new Set(srcFull.replace(/[.,;:!?'"()\\[\\]«»…–—‘’“”‹›°]/g,' ').split(/\\s+/).filter(w=>w.length>2&&!stopWords.has(w)));
        let w2 = new Set(tgtFull.replace(/[.,;:!?'"()\\[\\]«»…–—‘’“”‹›°]/g,' ').split(/\\s+/).filter(w=>w.length>2&&!stopWords.has(w)));
        let common = new Set();
        w1.forEach(w=>{{if(w2.has(w)) common.add(w)}});
        return common;
    }}

    function markCW(txt,cwSet){{
        if(!cwSet || cwSet.size === 0) return txt;
        return txt.replace(/\\S+/g, function(w){{
            let lw = w.toLowerCase().replace(/[.,;:!?'"()\\[\\]«»…–—‘’“”‹›°]/g,'');
            if(cwSet.has(lw)) return '<span class="cw">'+w+'</span>';
            return w;
        }});
    }}

    function hPanel(t,c,p,diffs,showDiff,showCW,diffClass,cwSet){{
        let s=p[0],x=p[1];
        let pr_raw = t.substring(Math.max(0,s-c),s);
        let suf_raw = t.substring(x,Math.min(t.length,x+c));
        let pr = e(pr_raw);
        let suf = e(suf_raw);
        let inner='';
        if(showDiff&&diffs&&diffs.length>0){{
            diffs.sort((a,b)=>a[0]-b[0]);
            let l=s;
            for(let j=0;j<diffs.length;j++){{
                let seg=e(t.substring(l,diffs[j][0]));
                if(showCW) seg=markCW(seg,cwSet);
                inner+='<span style="background:'+BG+';padding:1px 3px;border-radius:2px">'+seg+'</span>';
                inner+='<span class="'+diffClass+'">'+e(t.substring(diffs[j][0],diffs[j][1]))+'</span>';
                l=diffs[j][1];
            }}
            let last=e(t.substring(l,x));
            if(showCW) last=markCW(last,cwSet);
            inner+='<span style="background:'+BG+';padding:1px 3px;border-radius:2px">'+last+'</span>';
        }} else {{
            let seg=e(t.substring(s,x));
            if(showCW) seg=markCW(seg,cwSet);
            inner='<span style="background:'+BG+';padding:2px 4px;border-radius:3px">'+seg+'</span>';
        }}
        if(showCW){{
            pr = markCW(pr, cwSet);
            suf = markCW(suf, cwSet);
        }}
        return pr+inner+suf;
    }}

    function u(i){{
        let ss=document.getElementById('chk_sup_'+i).checked;
        let si=document.getElementById('chk_ins_'+i).checked;
        let cw=document.getElementById('chk_cw_'+i).checked;
        let cwSet=cw?getBlockCW(i):new Set();
        document.getElementById('a'+i).innerHTML=hPanel(t1,cs[i],m[i][0],m[i][2],ss,cw,'sup',cwSet);
        document.getElementById('b'+i).innerHTML=hPanel(t2,cs[i],m[i][1],m[i][3],si,cw,'ins',cwSet);
    }}

    function c(i,d){{cs[i]=Math.max({ctx},cs[i]+d);u(i)}}
    function r(i){{cs[i]={ctx};u(i)}}

    // ── Gestion de la sélection ────────────────────────────────────────────
    function updateSelCount(){{
        let total=selectedGlobal.size;
        let label=total>0&&total===allM.length
            ? total+' sélectionné(s) — toutes les pages'
            : total+' sélectionné(s)';
        document.getElementById('sel-count').textContent=label;
        let btn=document.getElementById('sel-all-btn');
        if(btn) btn.textContent=allPagesSelected?'☐ Tout désélectionner':'☑ Tout sélectionner';
    }}

    function onCheckChange(localIdx){{
        let globalIdx=startOffset+localIdx;
        let el=document.getElementById('chk_sel_'+localIdx);
        if(el&&el.checked){{
            selectedGlobal.add(globalIdx);
        }} else {{
            selectedGlobal.delete(globalIdx);
            allPagesSelected=false;
        }}
        updateSelCount();
    }}

    function toggleAll(){{
        allPagesSelected=!allPagesSelected;
        if(allPagesSelected){{
            for(let i=0;i<allM.length;i++) selectedGlobal.add(i);
            for(let i=0;i<m.length;i++){{
                let el=document.getElementById('chk_sel_'+i);
                if(el) el.checked=true;
            }}
        }} else {{
            selectedGlobal.clear();
            for(let i=0;i<m.length;i++){{
                let el=document.getElementById('chk_sel_'+i);
                if(el) el.checked=false;
            }}
        }}
        updateSelCount();
    }}

    function exportCSV(){{
        if(selectedGlobal.size===0){{alert('Sélectionnez au moins un bloc à exporter.');return;}}
        let rows=[['num','pos_source','pos_cible','texte_source','texte_cible','nb_suppressions','nb_insertions']];
        let sortedSel=Array.from(selectedGlobal).sort((a,b)=>a-b);
        for(let k=0;k<sortedSel.length;k++){{
            let gi=sortedSel[k];
            let item=allM[gi];
            let p1=item[0],p2=item[1];
            let ns=item[2],ni=item[3];
            let src=t1.substring(p1[0],p1[1]).substring(0,200);
            let tgt=t2.substring(p2[0],p2[1]).substring(0,200);
            rows.push([(gi+1),p1[0]+'-'+p1[1],p2[0]+'-'+p2[1],'"'+src.replace(/"/g,'""')+'"','"'+tgt.replace(/"/g,'""')+'"',ns,ni]);
        }}
        let csv=rows.map(r=>r.join(',')).join('\\n');
        let blob=new Blob([csv],{{type:'text/csv;charset=utf-8;'}});
        let url=URL.createObjectURL(blob);
        let a=document.createElement('a');a.href=url;a.download='alignements_selection.csv';a.click();
        URL.revokeObjectURL(url);
    }}

    for(let i=0;i<m.length;i++) u(i);
    for(let i=0;i<m.length;i++){{
        let el=document.getElementById('chk_sel_'+i);
        if(el&&selectedGlobal.has(startOffset+i)) el.checked=true;
    }}
    updateSelCount();
    </script></body></html>"""
    
    return html


if __name__ == "__main__":
    main()